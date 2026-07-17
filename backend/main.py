"""
Bajaj Life Insurance – Speech Analytics Platform
FastAPI Backend
"""

import os, json, uuid, re, asyncio, hashlib, shutil, io
from collections import Counter, defaultdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, List, Dict, Any
import logging

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Query
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

import pandas as pd
import openpyxl
import numpy as np
from pypdf import PdfReader
from docx import Document as DocxDocument

import openai
from dotenv import load_dotenv
import uuid as _uuid_mod  # already imported as uuid above — no change needed
 
from backend.db import (
    create_job, update_job, get_job, list_jobs, increment_job_counters,
    insert_raw_call, get_raw_call,
    upsert_analyzed_call, get_analyzed_call, delete_analyzed_call,
    list_analyzed_calls, get_fatal_calls as db_get_fatal_calls,
    get_dashboard_stats, clear_all_data,
)
try:
    import faiss
except Exception:  # pragma: no cover - optional dependency
    faiss = None

# ── Config ──────────────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent.parent
UPLOAD_DIR = BASE_DIR / "uploads"
PRODUCT_DIR= BASE_DIR / "products"
PROC_DIR   = BASE_DIR / "processed"
DB_FILE    = PROC_DIR / "calls_db.json"
CHROMA_DIR = BASE_DIR / "chroma_db"
PRODUCT_INDEX_FILE = PROC_DIR / "product_specs_index.json"
RAG_INDEX_FILE = PROC_DIR / "product_faiss.index"
RAG_META_FILE = PROC_DIR / "product_faiss_meta.json"
RAG_BACKEND_FILE = PROC_DIR / "rag_backend.json"
RAG_EMBED_MODEL = os.environ.get("RAG_EMBED_MODEL", "all-MiniLM-L6-v2").strip()
PRODUCT_SOURCES = [PRODUCT_DIR]
os.environ["HF_HUB_DISABLE_SSL_VERIFICATION"] = "1"
for d in [UPLOAD_DIR, PRODUCT_DIR, PROC_DIR, CHROMA_DIR]:
    d.mkdir(parents=True, exist_ok=True)

load_dotenv(BASE_DIR / ".env")

def _env(name: str, default: str = "") -> str:
    return os.environ.get(name, default).strip()

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("speech_analytics")

# OpenAI async client (supports OpenAI and Azure OpenAI)
# Using AsyncAzureOpenAI / AsyncOpenAI so that chat.completions.create() can be
# properly awaited without blocking the event loop.
if _env("AZURE_OPENAI_API_KEY") and _env("AZURE_OPENAI_ENDPOINT"):
    OPENAI_MODEL = _env("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")
    openai_client = openai.AsyncAzureOpenAI(
        api_key=_env("AZURE_OPENAI_API_KEY"),
        api_version=_env("AZURE_OPENAI_API_VERSION", "2024-02-01"),
        azure_endpoint=_env("AZURE_OPENAI_ENDPOINT")
    )
else:
    OPENAI_MODEL = _env("OPENAI_MODEL", "gpt-4o")
    openai_client = openai.AsyncOpenAI(api_key=_env("OPENAI_API_KEY", "YOUR_OPENAI_API_KEY"))

# ── FastAPI App ──────────────────────────────────────────────────────────────
app = FastAPI(title="Bajaj Life Insurance – Speech Analytics", version="1.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# Serve frontend
FRONTEND_DIR = BASE_DIR / "frontend"
app.mount("/static", StaticFiles(directory=str(FRONTEND_DIR / "static")), name="static")

# ── RAG Setup ────────────────────────────────────────────────────────────────
rag_collection = None

def init_rag():
    """Initialize the local product knowledge base."""
    global rag_collection
    rag_collection = None
    try:
        rebuild_product_rag_index()
        log.info("Product RAG index initialized")
    except Exception as e:
        log.warning(f"RAG init failed (non-fatal): {e}")

# @app.on_event("startup")
# async def startup():
#     init_rag()
#     clear_expired_cache()  # Clean up any expired cache files
#     if not DB_FILE.exists():
#         DB_FILE.write_text(json.dumps({"calls": [], "jobs": []}, indent=2))
#     if not PRODUCT_INDEX_FILE.exists():
#         PRODUCT_INDEX_FILE.write_text(json.dumps({"products": []}, indent=2))
@app.on_event("startup")
async def startup():
    init_rag()
    clear_expired_cache()
    if not PRODUCT_INDEX_FILE.exists():
        PRODUCT_INDEX_FILE.write_text(json.dumps({"products": []}, indent=2))
    # Verify DB connection on startup (PostgreSQL mode only)
    from backend.db import DATABASE_URL as _DB_URL
    if _DB_URL:
        try:
            from backend.db import get_conn
            conn = get_conn()
            conn.close()
            log.info("PostgreSQL connection OK")
        except Exception as e:
            log.error(f"PostgreSQL connection FAILED: {e}")
    else:
        log.info("Running in JSON-file mode (no DATABASE_URL configured)")

# ── Database helpers ─────────────────────────────────────────────────────────
def load_db() -> dict:
    try:
        return json.loads(DB_FILE.read_text())
    except Exception:
        return {"calls": [], "jobs": []}

def save_db(db: dict):
    DB_FILE.write_text(json.dumps(db, indent=2, default=str))

def load_product_index() -> dict:
    try:
        return json.loads(PRODUCT_INDEX_FILE.read_text())
    except Exception:
        return {"products": []}

def save_product_index(index_data: dict):
    PRODUCT_INDEX_FILE.write_text(json.dumps(index_data, indent=2, default=str))

USAGE_FILE = PROC_DIR / "api_usage_log.json"

def log_api_usage(model: str, input_tokens: int, output_tokens: int):
    entry = {"timestamp": datetime.now().isoformat(), "model": model,
             "input_tokens": input_tokens, "output_tokens": output_tokens}
    try:
        existing = json.loads(USAGE_FILE.read_text()) if USAGE_FILE.exists() else []
        existing.append(entry)
        USAGE_FILE.write_text(json.dumps(existing, indent=2))
    except Exception as e:
        log.warning(f"Usage log error: {e}")

# ── Job Control ───────────────────────────────────────────────────────────────
# In-memory store for cancellation/pause signals.
# Keys are job_ids, values are "cancel" | "pause" | "resume"
_job_control: Dict[str, str] = {}

def signal_job(job_id: str, signal: str):
    """Send a control signal to a running job."""
    _job_control[job_id] = signal

def check_job_signal(job_id: str) -> Optional[str]:
    """Check if a job has a pending control signal. Clears it after reading."""
    return _job_control.get(job_id)

def clear_job_signal(job_id: str):
    _job_control.pop(job_id, None)

# ── Cache System (long-lived: expires in 2+ weeks) ──────────────────────────
CACHE_DIR = PROC_DIR / "cache"
CACHE_DIR.mkdir(parents=True, exist_ok=True)
CACHE_EXPIRY_SECONDS = 14 * 24 * 60 * 60  # 2 weeks

def _get_cache_file(key: str) -> Path:
    """Get cache file path for a given key"""
    safe_key = hashlib.md5(key.encode()).hexdigest()
    return CACHE_DIR / f"{safe_key}.json"

def get_cache(key: str) -> Any:
    """Get value from cache if not expired"""
    cache_file = _get_cache_file(key)
    if not cache_file.exists():
        return None
    try:
        data = json.loads(cache_file.read_text())
        age_seconds = (datetime.now() - datetime.fromisoformat(data.get("cached_at", ""))).total_seconds()
        if age_seconds < CACHE_EXPIRY_SECONDS:
            return data.get("value")
        else:
            cache_file.unlink(missing_ok=True)
    except Exception:
        cache_file.unlink(missing_ok=True)
    return None

def set_cache(key: str, value: Any):
    """Store value in cache with timestamp"""
    cache_file = _get_cache_file(key)
    try:
        cache_file.write_text(json.dumps({
            "cached_at": datetime.now().isoformat(),
            "value": value
        }, default=str))
    except Exception as e:
        log.warning(f"Cache write error: {e}")

def clear_expired_cache():
    """Remove expired cache files (called periodically)"""
    try:
        for cache_file in CACHE_DIR.glob("*.json"):
            try:
                data = json.loads(cache_file.read_text())
                age_seconds = (datetime.now() - datetime.fromisoformat(data.get("cached_at", ""))).total_seconds()
                if age_seconds >= CACHE_EXPIRY_SECONDS:
                    cache_file.unlink(missing_ok=True)
            except Exception:
                cache_file.unlink(missing_ok=True)
    except Exception as e:
        log.warning(f"Cache cleanup error: {e}")


def _load_embedder():
    if hasattr(_load_embedder, "_model"):
        return _load_embedder._model
    try:
        from sentence_transformers import SentenceTransformer
        _load_embedder._model = SentenceTransformer(RAG_EMBED_MODEL)
    except Exception:
        _load_embedder._model = None
    return _load_embedder._model


def _hash_vector(text: str, dim: int = 384) -> np.ndarray:
    vec = np.zeros(dim, dtype=np.float32)
    tokens = re.findall(r"[a-z0-9]+", (text or "").lower())
    for token in tokens:
        index = int(hashlib.md5(token.encode("utf-8")).hexdigest(), 16) % dim
        vec[index] += 1.0
    norm = float(np.linalg.norm(vec))
    if norm > 0:
        vec /= norm
    return vec


def _embed_texts(texts: List[str]) -> np.ndarray:
    model = _load_embedder()
    if model is not None:
        return model.encode(texts, convert_to_numpy=True, normalize_embeddings=True).astype(np.float32)
    return np.vstack([_hash_vector(text) for text in texts]).astype(np.float32)


def _safe_filename_label(name: str) -> str:
    stem = Path(name).stem
    stem = re.sub(r"(?i)product\s*circular", "", stem)
    stem = re.sub(r"[_\-]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    stem = re.sub(r"(?i)\bbajaj\s+allianz\s+life\b", "Bajaj Life", stem).strip()
    stem = re.sub(r"\s{2,}", " ", stem).strip(" -_")
    return stem or Path(name).stem

_PARAM_WEIGHTS = {
    "greeting_opening": 5,
    "query_understanding": 10,
    "response_accuracy": 25,
    "communication_quality": 8,
    "compliance": 20,
    "personalisation": 5,
    "empathy_soft_skills": 5,
    "resolution": 10,
    "system_behaviour": 10,
    "closing_interaction": 2,
}

_PARAM_MIN_PASS = {
    "greeting_opening": 3,
    "query_understanding": 3,
    "response_accuracy": 4,
    "communication_quality": 3,
    "compliance": 4,
    "personalisation": 3,
    "empathy_soft_skills": 3,
    "resolution": 3,
    "system_behaviour": 3,
    "closing_interaction": 3,
}

def _compute_weighted_score(scores: Dict[str, Any]) -> float:
    """Always compute server-side — never trust the model's weighted_score."""
    total_weight = sum(_PARAM_WEIGHTS.values())  # 100
    weighted = sum(
        float(scores.get(k, 3) or 3) * w
        for k, w in _PARAM_WEIGHTS.items()
    )
    # Scores are 1-5, weights sum to 100, so max = 5*100 = 500
    # Normalise to 0-100
    return round(weighted / 5.0 * (100 / total_weight), 2)

def _compute_failed_parameters(scores: Dict[str, Any]) -> List[str]:
    return [k for k, min_v in _PARAM_MIN_PASS.items() if int(scores.get(k, 0) or 0) < min_v]
PARAM_ORDER = [
    "greeting_opening", "query_understanding", "response_accuracy", "communication_quality",
    "compliance", "personalisation", "empathy_soft_skills", "resolution",
    "system_behaviour", "closing_interaction"
]


def _fallback_param_comments(scores: Dict[str, Any]) -> List[str]:
    text = {
        "greeting_opening": "Opening tone and welcome quality from the first bot turn.",
        "query_understanding": "How accurately the bot interpreted customer intent and follow-up questions.",
        "response_accuracy": "Factual correctness of product/policy information provided by the bot; wrong product facts are fatal.",
        "communication_quality": "Clarity, structure, and readability of the bot's language.",
        "compliance": "Compliance with privacy, regulatory constraints, and staying on-product for the question asked.",
        "personalisation": "Use of customer context, policy context, and personalization cues.",
        "empathy_soft_skills": "Warmth, reassurance, and acknowledgement of customer concerns.",
        "resolution": "Whether the customer issue was actually solved or moved forward.",
        "system_behaviour": "Flow quality: no loops, no instability, and no repetitive failures.",
        "closing_interaction": "Quality and completeness of the call closure.",
    }
    out = []
    for key in PARAM_ORDER:
        score = int(scores.get(key, 3) or 3)
        suffix = "Strong evidence." if score >= 4 else ("Acceptable but with gaps." if score == 3 else "Needs corrective action.")
        out.append(f"{text[key]} {suffix}")
    return out


def _score_reason(analysis: Dict[str, Any]) -> str:
    failed = analysis.get("failed_parameters") or []
    if failed:
        return f"Failed parameters: {', '.join(failed)}.".strip()
    return "Weighted score computed from parameter scores and policy thresholds."


def _refine_sentiment(turns: List[Dict], model_sentiment: str = "neutral") -> str:
    """Tie-breaker only: trust the model unless it returned an unrecognised value,
    or it returned plain 'neutral' AND the customer text contains strong, unambiguous
    emotional signals (full phrases, not isolated keywords like 'still' / 'issue')."""
    allowed = {"positive", "neutral", "frustrated", "angry", "distressed"}
    model_sentiment = (model_sentiment or "neutral").strip().lower()

    # If model returned something non-neutral and valid, ALWAYS keep it.
    if model_sentiment in allowed and model_sentiment != "neutral":
        return model_sentiment

    # Only override 'neutral' if there is strong evidence (multi-word phrases).
    customer_text = " ".join(
        t.get("text", "") for t in turns if (t.get("speaker") or "").lower() in {"customer", "user"}
    ).lower()

    # Strong, unambiguous distress (whole phrases)
    if re.search(r"\b(this is urgent|emergency|i('?| a)m scared|i('?| a)m panicking|please help me)\b", customer_text):
        return "distressed"
    # Explicit anger (whole phrases or strong single words)
    if re.search(r"\b(this is unacceptable|absolutely ridiculous|i('?| a)m (very )?angry|worst service|terrible service|extremely frustrated)\b", customer_text):
        return "angry"
    # Explicit positive sentiment (gratitude phrases)
    if re.search(r"\b(thank you so much|that was very helpful|really appreciate|excellent service|perfectly resolved)\b", customer_text):
        return "positive"

    return model_sentiment if model_sentiment in allowed else "neutral"


def _annotate_transcript(turns: List[Dict], analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
    annotations = []
    sentiment = (analysis.get("sentiment") or "neutral").lower()
    flags = set(analysis.get("flags") or [])

    for idx, turn in enumerate(turns, start=1):
        speaker = turn.get("speaker", "unknown")
        text = (turn.get("text") or "").strip()
        text_lower = text.lower()
        tags = []

        if speaker in {"bot", "agent", "system"}:
            if re.search(r"grace period|30 day", text, re.I):
                tags.append("Grace Period Explained")
            if re.search(r"transfer|connect you to|handover|escalate", text, re.I):
                tags.append("Escalation Action")
            if re.search(r"thank you|thanks|happy to help|is there anything else", text, re.I):
                tags.append("Closure Cue")
            if re.search(r"premium|sum assured|maturity|policy term|coverage|death benefit|rider", text, re.I):
                tags.append("Product Information")
            if re.search(r"sorry|apologize|apologies|regret|understand your concern", text, re.I):
                tags.append("Empathy Shown")
            if re.search(r"as per|according to|our records show|our system|we can see", text, re.I):
                tags.append("Data Referenced")
            if re.search(r"unable to|cannot provide|can't provide|not able to", text, re.I):
                tags.append("⚠ Unable to Respond")
            if re.search(r"please note|important|kindly note|please be aware", text, re.I):
                tags.append("Important Notice")
            if re.search(r"surrender|loan against|partial withdrawal|discontinue", text, re.I):
                tags.append("Policy Action")
            if re.search(r"verified|confirmed|otp|authentication", text, re.I):
                tags.append("Verification")

        if speaker in {"customer", "user"}:
            if re.search(r"angry|frustrated|unacceptable|worst|ridiculous|complaint", text, re.I):
                tags.append("😠 Customer Frustrated")
            if re.search(r"urgent|emergency|panic|distress|scared", text, re.I):
                tags.append("🆘 Urgent")
            if re.search(r"thank|thanks|great|helpful|good|excellent", text, re.I):
                tags.append("😊 Positive Feedback")
            if re.search(r"why|how|what|when|where|can you|could you|please explain", text, re.I):
                tags.append("❓ Question")
            if re.search(r"claim|surrender|cancel|withdraw|close|terminate", text, re.I):
                tags.append("Policy Request")

        annotations.append({
            "sl": turn.get("sl", idx),
            "speaker": speaker,
            "text": text,
            "tags": tags,
        })
    return annotations


def _build_qa_findings(analysis: Dict[str, Any]) -> List[Dict[str, str]]:
    findings = []
    for fp in analysis.get("failed_parameters", [])[:6]:
        findings.append({"type": "negative", "text": f"Parameter below threshold: {fp.replace('_', ' ')}"})
    for flag in analysis.get("flags", [])[:6]:
        findings.append({"type": "warning", "text": f"Flag detected: {flag.replace('_', ' ')}"})
    issues = (analysis.get("product_issues") or "").strip()
    if issues and issues.lower() != "none":
        findings.append({"type": "warning", "text": f"Product check issue: {issues}"})
    strengths = (analysis.get("strengths") or "").strip()
    if strengths and strengths.lower() not in {"none", "n/a"}:
        findings.append({"type": "positive", "text": strengths})
    if not findings:
        findings.append({"type": "positive", "text": "No major QA issues detected in automated analysis."})
    return findings[:8]


def _sentence_split(text: str) -> List[str]:
    return [part.strip() for part in re.split(r"(?<=[.!?])\s+", text or "") if part.strip()]


def _normalize_number_token(token: str) -> str:
    token = (token or "").strip().lower().rstrip("%")
    try:
        value = float(token)
        return str(int(value)) if value.is_integer() else f"{value:.2f}".rstrip("0").rstrip(".")
    except Exception:
        return token


def _extract_number_tokens(text: str) -> List[str]:
    return [_normalize_number_token(token) for token in re.findall(r"\d+(?:\.\d+)?%?", text or "")]


def _bot_sentences(turns: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    sentences: List[Dict[str, Any]] = []
    for turn in turns or []:
        speaker = (turn.get("speaker") or "").lower()
        if speaker not in {"bot", "agent", "system"}:
            continue
        for sentence in _sentence_split(turn.get("text", "")):
            sentences.append({"sl": turn.get("sl"), "speaker": speaker, "text": sentence})
    return sentences


def _best_spec_sentence(statement: str, analysis: Dict[str, Any], keywords: List[str]) -> str:
    """Find the most supportive sentence for `statement` in product evidence.

    When the analysis contains multiple products (primary + secondaries from
    `infer_product_context`), restrict search to the rows whose `product`
    matches the bot's currently identified product. This stops comparisons
    against the wrong product spec from generating bogus 'risk' verdicts.
    """
    primary_product = (analysis.get("product_mentioned") or "").strip().lower()
    secondary = analysis.get("secondary_products") or []
    candidate_products = {primary_product}
    for s in secondary:
        p = (s.get("product") or "").strip().lower()
        if p:
            candidate_products.add(p)

    search_rows = list(analysis.get("product_evidence") or [])
    if not search_rows:
        search_rows = _load_rag_meta()
    profile = analysis.get("product_profile") or {}

    # Build the search text pool, biased toward the right product.
    def _row_product(r: Dict[str, Any]) -> str:
        return (r.get("product") or "").strip().lower()

    if primary_product and primary_product != "none":
        primary_texts = [r.get("text", "") for r in search_rows if _row_product(r) == primary_product and r.get("text")]
        other_texts = [r.get("text", "") for r in search_rows if _row_product(r) != primary_product and r.get("text")]
        # Search the right product first; fall back to others only if nothing scored
        search_texts = primary_texts or other_texts
    else:
        search_texts = [r.get("text", "") for r in search_rows if r.get("text")]

    search_texts.extend(profile.get("evidence_snippets") or [])
    if profile.get("summary"):
        search_texts.append(profile.get("summary"))

    stmt_tokens = set(_tokenize(statement))
    candidates: List[tuple] = []
    for text in search_texts:
        for sentence in _sentence_split(text):
            sentence_lower = sentence.lower()
            overlap = len(stmt_tokens.intersection(_tokenize(sentence)))
            keyword_hits = sum(1 for keyword in keywords if keyword in sentence_lower)
            score = overlap + (2 * keyword_hits)
            if score > 0:
                candidates.append((score, sentence.strip()))

    if not candidates:
        return (profile.get("summary") or "").strip()[:260]

    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _fact_contains_range(fact: str) -> bool:
    """Heuristic: does the fact sentence describe a numeric range (so a single
    value inside it is NOT a contradiction)? Avoids false-positive 'fail'
    verdicts when the bot states a value that falls within a spec range."""
    f = (fact or "").lower()
    if re.search(r"\b(between|from)\s+[\d.]+\s*(to|-|and|–|—)\s*[\d.]+", f):
        return True
    if re.search(r"\b(up to|upto|min(imum)?|max(imum)?|at least|at most|starting (from|at))\b", f):
        return True
    if re.search(r"[\d.]+\s*(to|-|–|—)\s*[\d.]+", f):  # bare "1 to 50", "18-65"
        return True
    return False


def _fact_is_meta_instruction(fact: str) -> bool:
    """Detect product-doc sentences that are meta-instructions to the agent
    (e.g. "verify from the system", "refer to the policy bond") rather than
    statements of actual product values. These must NEVER be used to
    contradict a bot statement — they are guidance, not facts."""
    if not fact:
        return False
    f = fact.lower()
    patterns = [
        r"verif(y|ied) (from|with|against) (the )?system",
        r"refer to (the )?(policy|product) (bond|document|brochure|circular)",
        r"check (the )?(system|records|cms|crm)",
        r"as per (the )?(records|system|customer|policyholder)",
        r"to be confirmed",
        r"please (verify|confirm|check)",
        r"should be verified",
        r"available in (the )?(system|customer master|policy schedule)",
    ]
    return any(re.search(p, f) for p in patterns)


def _customer_disputed_bot(turns: List[Dict[str, Any]]) -> bool:
    """Did the customer push back on something the bot said?
    Used to gate false_information flagging: we only mark a bot statement as
    false if the customer disputes it OR the spec clearly contradicts it.
    A bot quoting a number is NOT inaccurate on its own."""
    dispute_patterns = re.compile(
        r"\b("
        r"that('?s| is) (wrong|incorrect|not right|not correct)|"
        r"you (are|'re) wrong|"
        r"that('?s)? not (true|right|what i|correct)|"
        r"no(,| ) (that|it)('?s| is)? (not|wrong)|"
        r"incorrect|mistake|wrongly|misinform"
        r")\b",
        re.IGNORECASE,
    )
    for i, t in enumerate(turns or []):
        speaker = (t.get("speaker") or "").lower()
        if speaker not in {"customer", "user"}:
            continue
        text = (t.get("text") or "").strip()
        if not text:
            continue
        if dispute_patterns.search(text):
            # Make sure there was a bot turn just before — i.e. customer is
            # disputing the bot, not making a standalone complaint.
            for j in range(i - 1, max(-1, i - 3), -1):
                prev_speaker = (turns[j].get("speaker") or "").lower()
                if prev_speaker in {"bot", "agent", "system"}:
                    return True
    return False


# ── Customer-specific vs product-spec discriminator ─────────────────────────
# When the bot reads back data fetched from the CRM / policy administration
# system (the customer's own policy name, sum assured, maturity date, premium
# amount, nominee, etc.) it is NOT making a product claim — it is quoting
# that customer's record. Comparing such statements against the product spec
# is meaningless and generates false-positive 'inaccurate' verdicts. These
# helpers identify those statements so the QA rule layer can skip them.

_POSSESSIVE_CRM_ATTRS = (
    r"policy(?:\s+name|\s+number|\s+term|\s+holder)?|"
    r"sum\s+assured|premium(?:\s+amount)?|nominee|"
    r"maturity(?:\s+date|\s+amount|\s+value|\s+benefit)?|"
    r"surrender(?:\s+date|\s+value|\s+amount)?|"
    r"fund\s+value|account(?:\s+number)?|customer(?:\s+id)?|client(?:\s+id)?|"
    r"date\s+of\s+(?:birth|maturity|commencement|inception|issue)|"
    r"address|phone|mobile|email|contact|registered"
)

_POSSESSIVE_PATTERN = re.compile(
    rf"\b(?:your|the\s+customer'?s|this\s+customer'?s|policy\s*holder'?s|their|his|her)\s+"
    rf"(?:{_POSSESSIVE_CRM_ATTRS})\b",
    re.IGNORECASE,
)

_CRM_INTRODUCERS = (
    "your policy name is",
    "your policy is",
    "your policy will",
    "your sum assured",
    "your maturity",
    "your premium",
    "your nominee",
    "your surrender",
    "your fund value",
    "your policy term",
    "your date of",
    "your address",
    "your contact",
    "your registered",
    "as per our records, your",
    "as per your policy",
    "as per the records,",
    "as per our records,",
    "according to our records",
    "according to the system",
    "according to your policy",
    "our records show that your",
    "our records indicate",
    "i can see that your",
    "i can see your policy",
    "i can see your",
    "we can see that your",
    "policy is currently",
    "policy is in",
    "policy is active",
    "policy is surrendered",
    "policy has been",
)


def _is_customer_specific_statement(stmt: str) -> bool:
    """Detect if a bot statement is reading back CUSTOMER-SPECIFIC data from CRM.

    Such statements (e.g. "your policy name is Bajaj Life Goal Assure",
    "the sum assured is rupees seven lakh", "your maturity date is 22nd July
    2029") MUST NOT be compared against product specs because the spec only
    describes the product's structure / ranges, while the bot is quoting the
    customer's individual record.

    Returns True when the statement is clearly a CRM read-back.
    """
    if not stmt:
        return False
    s = stmt.lower().strip()

    # 1) Possessive customer reference + policy attribute → CRM read-back
    if _POSSESSIVE_PATTERN.search(s):
        return True

    # 2) Customer-data introducer phrases
    if any(intro in s for intro in _CRM_INTRODUCERS):
        return True

    # 3) Specific date attached to a policy/maturity/surrender context →
    #    almost certainly a per-customer date, not a product fact.
    has_specific_date = bool(
        re.search(
            r"\b\d{1,2}(?:st|nd|rd|th)?\s+(?:of\s+)?"
            r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
            r"jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
            r"[a-z]*[, ]+(?:twenty\s+)?\d{2,4}",
            s,
        )
        or re.search(r"\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b", s)
    )
    if has_specific_date and re.search(
        r"\b(your|policy|maturity|surrender|commenc(?:e|ed|ement)|"
        r"inception|effective|issue|premium\s+due|next\s+premium|paid)\b",
        s,
    ):
        return True

    # 4) Currency amount + possessive/policy framing → customer-specific
    #    e.g. "the sum assured is rupees seven lakh twenty thousand"
    has_currency_amount = bool(
        re.search(
            r"\b(?:rupees|rs\.?|inr|₹)\s*[\w\s,]{0,40}?"
            r"(?:lakh|crore|thousand|hundred|\d{3,})",
            s,
        )
    )
    if has_currency_amount and re.search(
        r"\b(?:your|sum\s+assured|premium|maturity|fund\s+value|surrender)\b", s
    ):
        return True

    return False


# ── System-failure detector ─────────────────────────────────────────────────
# Broader than the original "unable to provide information" check — catches
# any phrasing where the bot says it could not fetch / retrieve / access
# customer or policy data. These responses are HARD system failures and must
# be flagged as fatal (not "conditionally OK") because the customer left the
# call without the data they asked for.

_SYSTEM_FAILURE_PATTERNS = (
    # "unable to / cannot / couldn't <verb>" verbs that imply data retrieval
    re.compile(
        r"\b(?:unable\s+to|cannot|can'?t|could\s*not|couldn'?t|failed\s+to|"
        r"not\s+able\s+to|having\s+trouble\s+(?:to\s+)?)\s+"
        r"(?:provide|fetch|retriev\w*|access|find|locate|get|pull|obtain|"
        r"share|give|display|show|determine|confirm|verify|process|"
        r"check|look\s*up|generate)\b",
        re.IGNORECASE,
    ),
    # "I'm unable / not able / having difficulty" without explicit verb
    re.compile(
        r"\b(?:i'?m|i\s+am|we'?re|we\s+are|the\s+system\s+is)\s+"
        r"(?:unable|not\s+able|having\s+(?:difficulty|trouble|issues?))\b",
        re.IGNORECASE,
    ),
    # "Unfortunately, I couldn't / I can't / I am unable / we are unable"
    re.compile(
        r"\bunfortunately[,\s]+(?:i\s+(?:couldn'?t|could\s+not|can(?:not|'t)|"
        r"am\s+unable|was\s+unable)|we\s+(?:are\s+)?(?:unable|cannot|can'?t))\b",
        re.IGNORECASE,
    ),
    # System / technical issues
    re.compile(
        r"\b(?:facing\s+(?:an?\s+)?(?:issue|problem|error|outage)|"
        r"system\s+(?:error|issue|unavailable|down|outage|glitch)|"
        r"technical\s+(?:issue|error|difficulty|problem|glitch)|"
        r"experiencing\s+(?:an?\s+)?(?:issue|problem|error|delay))\b",
        re.IGNORECASE,
    ),
    # "right now" / "at this moment" / "currently" combined with inability
    re.compile(
        r"\b(?:right\s+now|at\s+(?:this|the)\s+moment|at\s+this\s+time|currently)\b"
        r"[^.?!]{0,80}?\b(?:unavailable|not\s+available|unable|cannot|"
        r"can'?t|couldn'?t|could\s*not|won'?t\s+be\s+able)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:unable|cannot|can'?t|couldn'?t|could\s*not)\b"
        r"[^.?!]{0,80}?\b(?:right\s+now|at\s+(?:this|the)\s+moment|"
        r"at\s+this\s+time|currently|temporarily)\b",
        re.IGNORECASE,
    ),
)


def _is_system_failure_response(text: str) -> bool:
    """Detect any bot response indicating it failed to fetch / retrieve /
    access requested data — broader than the legacy 'unable to provide'
    phrase set. Returns True when the bot has effectively told the customer
    the data they asked for is currently unavailable."""
    if not text:
        return False
    return any(pat.search(text) for pat in _SYSTEM_FAILURE_PATTERNS)

_ESCALATION_PATTERNS = re.compile(
    r"\b(?:transfer(?:ring)?\s+you|connect(?:ing)?\s+you\s+to|"
    r"escalat(?:e|ing|ed)|hand(?:ing)?\s+over\s+to|"
    r"raise\s+a\s+(?:ticket|request)|loop\s+in\s+(?:an?\s+)?(?:agent|specialist)|"
    r"(?:an?\s+)?(?:agent|specialist|representative|relationship\s+manager)\s+"
    r"will\s+(?:get\s+back|reach\s+out|contact|call)|"
    r"schedul(?:e|ing)\s+a\s+callback)\b",
    re.IGNORECASE,
)

def _bot_escalated(turns: List[Dict[str, Any]]) -> bool:
    """True if any bot turn offers escalation/transfer/callback."""
    for t in turns or []:
        if (t.get("speaker") or "").lower() in {"bot", "agent", "system"}:
            if _ESCALATION_PATTERNS.search(t.get("text") or ""):
                return True
    return False

# ── Transcript-level product-name extractor ─────────────────────────────────
# Not all products mentioned in a call exist in the RAG index (e.g. a
# customer references "Bajaj Life Shield Insurance Plan" but only Goal
# Suraksha specs are uploaded). The RAG-based detector therefore misses
# those names. This regex extractor surfaces them straight from the
# transcript so the UI / multi-product list always reflects what was said.

_PRODUCT_NAME_PATTERNS = (
    # "Bajaj [Allianz] [Life] <Capitalised words> <Suffix>"
    re.compile(
        r"\bBajaj(?:\s+Allianz)?(?:\s+Life)?(?:\s+Insurance)?"
        r"(?:\s+[A-Z][A-Za-z&]+){1,5}"
        r"(?:\s+(?:Plan|Policy|Insurance|Suraksha|Assure|Shield|Goal|Wealth|"
        r"Smart|Term|Protector|Saver|Income|Sampoorn|Sampada|Lifestyle|Protect|Invest|"
        r"Magnum|Future|Care|Eternal|Premier|Elite|Lakshya|Sanchay|Anand|ACE|Guaranteed|Benefit|Care|"
        r"Money|ULIP|eTouch|Family|Flexi|Future|Supreme|Endowment))?\b"
    ),
    # Stand-alone Bajaj product families ending with Plan/Policy/Insurance
    re.compile(
        r"\b(?:Goal|Smart|Lifestyle|Magnum|Future|Sampoorn|Sampada|Lakshya|"
        r"Sanchay|Anand|Eternal|Premier|eTouch|Guaranteed|Wealth|Cash|Shield|Wealth)\s+"
        r"(?:[A-Z][a-z]+\s+){0,3}"
        r"(?:Plan|Policy|Insurance|Suraksha|Assure|Shield)\b"
    ),
)
_GENERIC_PRODUCT_NAMES = {"none", "unknown", "", "bajaj life", "bajaj life insurance",
                           "bajaj allianz life", "bajaj allianz life insurance", "bajaj allianz", "bajaj"}


def _extract_product_mentions_from_text(text: str) -> List[str]:
    if not text:
        return []

    # Generic brand-only phrases that are NOT product names — must be rejected
    # even after cleaning/title-casing, since they appear in scripted openings
    # ("Bajaj Life Insurance mein aapka swagat hai") and would otherwise be
    # picked up as a false "explicit mention".
    _BRAND_ONLY = {
        "bajaj life", "bajaj life insurance", "bajaj allianz life",
        "bajaj allianz life insurance", "bajaj allianz", "bajaj",
    }

    matches: List[str] = []
    seen: set = set()
    for pattern in _PRODUCT_NAME_PATTERNS:
        for m in pattern.finditer(text):
            raw = re.sub(r"\s+", " ", m.group(0).strip())
            raw = re.sub(r"[\.,;:!\?]+$", "", raw).strip()
            if len(raw) < 6:
                continue
            words = []
            for w in raw.split():
                if w.lower() in {"of", "and", "the", "for", "with"}:
                    words.append(w.lower())
                elif w.isupper() and len(w) <= 4:
                    words.append(w)
                else:
                    words.append(w[:1].upper() + w[1:].lower() if w else w)
            cleaned = " ".join(words)
            cleaned = re.sub(r"(?i)\bbajaj\s+allianz\s+life\b", "Bajaj Life", cleaned)
            key = cleaned.lower()
            if key in _BRAND_ONLY:
                continue
            if key not in seen:
                seen.add(key)
                matches.append(cleaned)
    return matches


# ── Date-range filter helper (used by /api/calls, /api/fatal-calls, etc.) ──
def _filter_calls_by_range(
    calls: List[Dict[str, Any]], range_key: Optional[str]
) -> List[Dict[str, Any]]:
    """Filter call records by a friendly time range token.

    Accepted values (case-insensitive, several aliases per range):
      * 'all', 'forever', None, ''   → no filtering
      * 'week', '7d', 'last_week'    → last 7 days
      * 'month', '30d', 'last_month' → last 30 days
      * 'day', 'today', '24h'        → last 24 hours
    """
    if not range_key:
        return calls
    key = range_key.strip().lower()
    if key in {"", "all", "forever", "any"}:
        return calls
    now = datetime.now()
    if key in {"day", "today", "24h", "1d"}:
        cutoff = now - timedelta(days=1)
    elif key in {"week", "7d", "7days", "last_week", "lastweek"}:
        cutoff = now - timedelta(days=7)
    elif key in {"month", "30d", "30days", "last_month", "lastmonth"}:
        cutoff = now - timedelta(days=30)
    else:
        return calls
    cutoff_iso = cutoff.isoformat()
    out: List[Dict[str, Any]] = []
    for c in calls:
        ts = c.get("processed_at") or ""
        if ts >= cutoff_iso:
            out.append(c)
    return out


def _classify_product_check(statement: str, fact: str, keywords: List[str]) -> tuple:
    """Conservative rule-based verdict — only fire on hard, objective contradictions.

    Design:
      * Trust GPT-4o for semantic verdicts (it already produced its own product_checks).
      * Rule layer is ONLY a safety net for: (a) evasive bot responses on product
        questions, (b) clear numeric contradictions where the fact states a single
        specific value (not a range) that disagrees with the bot.
      * If fact contains a range (e.g. "1 to 50 lakhs"), single values are NOT
        flagged as contradictions — they may well be inside the range.
      * If the fact is a meta-instruction ("verify from system") it is NOT a
        product value — never use it to contradict the bot.
      * Default is 'pass' — never invent risk from token overlap.
    """
    stmt_lower = (statement or "").lower()
    fact_lower = (fact or "").lower()

    # NEW: customer-specific CRM read-backs are NEVER product-spec contradictions.
    # The spec describes the product's structure; the bot is quoting THIS
    # customer's individual data (their policy name, sum assured, maturity
    # date, etc.). Pass these unconditionally.
    if _is_customer_specific_statement(statement):
        return "pass", "None"

    if _is_system_failure_response(statement):
        return "fail", "🚨 HIGH — bot reported a system failure when asked for customer/product data"

    # NEW: meta-instructions in the spec are not facts. Skip.
    if _fact_is_meta_instruction(fact):
        return "pass", "None"

    stmt_nums = _extract_number_tokens(statement)
    fact_nums = _extract_number_tokens(fact)
    if not stmt_nums or not fact_nums:
        return "pass", "None"

    # Same numbers present → consistent.
    if set(stmt_nums).intersection(fact_nums):
        return "pass", "None"

    # Both sides must reference the same topic keyword to be comparable.
    stmt_has_topic = any(kw in stmt_lower for kw in keywords)
    fact_has_topic = any(kw in fact_lower for kw in keywords)
    if not (stmt_has_topic and fact_has_topic):
        return "pass", "None"

    # If the fact describes a range / min / max, a single statement value is
    # probably IN-range — do not call this a fail. Surface as a soft risk so
    # a human can confirm.
    if _fact_contains_range(fact):
        return "risk", "⚠️ Medium — bot stated a specific value where spec describes a range; please verify"

    # Both have specific single values, same topic, and they disagree → fail.
    return "fail", "🚨 HIGH — numeric detail conflicts with the product specification"


def _build_product_checks(turns: List[Dict[str, Any]], analysis: Dict[str, Any]) -> List[Dict[str, str]]:
    """Generate rule-based product checks ONLY for topics the model did not already
    cover. The model's product_checks are the primary source; this fills gaps for
    the unable_response + numeric-contradiction safety net.
    """
    product_name = str(analysis.get("product_mentioned") or "None").strip()
    if not product_name or product_name == "None":
        return []
        # ── NEW: skip if no matching spec exists for this product ────────────────
    # AFTER — use the full indexed catalog, not just this call's retrieved chunks
    _all_meta = _load_rag_meta()
    available_rag_products = {
        _safe_filename_label(row.get("product") or "").lower()
        for row in _all_meta
        if row.get("product")
    } if _all_meta else {
        # fallback: still use call evidence if meta file is missing
        _safe_filename_label(row.get("product") or "").lower()
        for row in (analysis.get("product_evidence") or [])
        if row.get("product")
    }
    product_lower = product_name.lower()
    # has_matching_spec = any(
    #     product_lower in p or p in product_lower
    #     or any(w in p for w in product_lower.split() if len(w) > 3)
    #     for p in available_rag_products
    # )
    # if not has_matching_spec:
    #     log.info(f"[QA] No matching RAG spec for '{product_name}' — skipping rule checks")
    #     return []
    def _word_overlap_ratio(a: str, b: str) -> float:
        wa = set(w for w in a.lower().split() if len(w) > 2)
        wb = set(w for w in b.lower().split() if len(w) > 2)
        if not wa or not wb:
            return 0.0
        return len(wa & wb) / min(len(wa), len(wb))

    has_matching_spec = any(
        _word_overlap_ratio(product_lower, p) >= 0.35
        for p in available_rag_products
    )
    if not has_matching_spec:
        log.info(f"[QA] No matching RAG spec for '{product_name}' (available: {list(available_rag_products)}) — skipping rule checks")
        return []
    # ── end guard ────────────────────────────────────────────────────────────

    # Skip topics the model has already produced a check for
    existing_topics = set()
    for chk in analysis.get("product_checks") or []:
        stmt_l = (chk.get("stmt") or "").lower()
        for topic_name, kws in [
            ("policy_term_eligibility", ["policy term", "entry age", "maturity age", "eligibility"]),
            ("premium_payment", ["premium", "monthly", "quarterly", "half-yearly", "yearly"]),
            ("benefits", ["maturity benefit", "death benefit", "guaranteed addition", "payout"]),
            ("surrender_loan", ["surrender", "loan", "paid-up", "grace period", "revival"]),
        ]:
            if any(kw in stmt_l for kw in kws):
                existing_topics.add(topic_name)

    bot_sentences = _bot_sentences(turns)
    topic_map = [
        ("policy_term_eligibility", ["policy term", "entry age", "maturity age", "eligibility", "age at maturity", "sum assured"]),
        ("premium_payment", ["premium", "monthly", "quarterly", "half-yearly", "half yearly", "yearly", "auto-debit", "modal premium"]),
        ("benefits", ["maturity benefit", "death benefit", "guaranteed additions", "sum assured", "payout", "cover"]),
        ("surrender_loan", ["surrender", "loan", "paid-up", "grace period", "revival", "foreclosure"]),
    ]

    checks: List[Dict[str, str]] = []
    for topic_name, keywords in topic_map:
        if topic_name in existing_topics:
            continue  # model already covered it
        statement_item = next(
            (item for item in bot_sentences if any(keyword in item["text"].lower() for keyword in keywords)),
            None,
        )
        if not statement_item:
            continue

        statement = statement_item["text"].strip()

        # Skip customer-specific (CRM read-back) statements outright — they
        # describe the customer's individual record, not the product structure,
        # and must not be compared against the spec.
        if _is_customer_specific_statement(statement):
            continue

        fact = _best_spec_sentence(statement, analysis, keywords)
        if not fact:
            continue

        verdict, risk = _classify_product_check(statement, fact, keywords)
        # Only emit non-pass verdicts from the rule layer to avoid spamming the UI
        # with low-value "pass" rows the model didn't bother generating.
        if verdict == "pass":
            continue

        checks.append({
            "call": str(analysis.get("call_id", "")),
            "stmt": statement,
            "fact": fact,
            "verdict": verdict,
            "vtext": "✓ ACCURATE" if verdict == "pass" else ("⚠️ CONDITIONALLY OK" if verdict == "risk" else "✗ INACCURATE"),
            "risk": risk,
            "topic": topic_name,
            "source": "rule",
        })

        if len(checks) >= 2:
            break

    return checks

def _restrict_failure_criteria(analysis: Dict[str, Any]) -> Dict[str, Any]:
    """Per business rule: a call should only FAIL if:
       (a) true system failure (no escalation offered), or
       (b) false_information (wrong product/policy facts), or
       (c) rude/poor communication (communication_quality below min pass).
    Other low scores (e.g. personalisation, closing) should not flip pass_fail."""
    flags = set(analysis.get("flags") or [])
    scores = analysis.get("scores") or {}
    comm_score = int(scores.get("communication_quality", 5) or 5)

    hard_fail_reasons = []
    if "system_failure" in flags:
        hard_fail_reasons.append("system_failure")
    if "false_information" in flags:
        hard_fail_reasons.append("false_information")
    if comm_score < _PARAM_MIN_PASS["communication_quality"]:
        hard_fail_reasons.append("communication_quality")

    analysis["pass_fail"] = "FAIL" if hard_fail_reasons else "PASS"
    analysis["fail_reasons"] = hard_fail_reasons
    # Keep failed_parameters for visibility/reporting, but it no longer drives pass_fail alone
    return analysis
def _apply_qa_policy_rules(analysis: Dict[str, Any], turns: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Safety-net rule layer.

    Design principles (do NOT re-derive verdicts via brittle keyword matching):
      1. TRUST the model's product_checks verdicts.
      2. Only generate rule-based checks for topics the model missed.
      3. To promote to FATAL we require BOTH:
           a. model produced (or rule produced) a 'fail' verdict, AND
           b. model returned a fatal-class flag (false_information / behavior_issue), AND
           c. EITHER two independent fail verdicts OR a 'fail' check whose 'risk'
              text contains a specific contradiction reason (numeric / spec / IRDAI / RBI).
         Otherwise the call is marked 'watch' — surfaced for human review, NOT auto-fatal.
      4. The unable_response + product_query → behavior_issue rule is preserved
         (it is unambiguous and the model often misses it).
      5. Flag list is deduped.
    """
    analysis = dict(analysis or {})
    scores = dict(analysis.get("scores") or {})
    flags = list(dict.fromkeys(analysis.get("flags") or []))
    failed_parameters = list(dict.fromkeys(analysis.get("failed_parameters") or []))

    # ------------------------------------------------------------------
    # Merge model checks with rule-generated checks (rule fills topic gaps only)
    # ------------------------------------------------------------------
    existing_checks = list(analysis.get("product_checks") or [])
    for c in existing_checks:
        c.setdefault("source", "model")
    generated_checks = _build_product_checks(turns, analysis)

    merged_checks: List[Dict[str, Any]] = []
    seen_pairs = set()
    for check in existing_checks + generated_checks:
        key = (
            (check.get("stmt") or "").strip().lower()[:160],
            (check.get("fact") or "").strip().lower()[:160],
        )
        if key in seen_pairs:
            continue
        seen_pairs.add(key)
        merged_checks.append(check)
    # analysis["product_checks"] = merged_checks[:6]
    # product_checks = analysis["product_checks"]
    # Filter out fail/risk checks that reference a different product's spec
    named_product = _safe_filename_label(
        str(analysis.get("product_mentioned") or "None")
    ).lower()
    filtered_checks: List[Dict[str, Any]] = []
    for c in merged_checks:
        check_product = _safe_filename_label(str(c.get("product") or "")).lower()
        verdict = (c.get("verdict") or "").lower()
        if check_product and check_product not in {"none", "unknown", ""}:
            same_product = (check_product in named_product) or (named_product in check_product)
            if not same_product and verdict in {"fail", "risk"}:
                log.info(
                    f"[QA] Dropping cross-product check: '{check_product}' "
                    f"vs call product '{named_product}'"
                )
                continue
        filtered_checks.append(c)

    analysis["product_checks"] = filtered_checks[:6]
    product_checks = analysis["product_checks"]

    # ------------------------------------------------------------------
    # Sanitise model verdicts:
    #   * If the bot statement is a customer-specific CRM read-back ("your
    #     policy name is X", "your maturity date is Y", "the sum assured is
    #     rupees Z"), downgrade ANY fail/risk → pass. The spec never describes
    #     a particular customer's policy, so these are not contradictions.
    #   * If the supporting "fact" is a meta-instruction ("verify from system",
    #     "refer to policy bond"), downgrade fail/risk → pass. The bot is not
    #     wrong just because the spec says "consult the system".
    #   * If verdict is fail/risk but no customer dispute happened AND no
    #     numeric contradiction is in the risk text, downgrade fail → risk
    #     (we surface for human review, not auto-fatal).
    # ------------------------------------------------------------------
    customer_disputed = _customer_disputed_bot(turns)
    for c in product_checks:
        verdict = (c.get("verdict") or "").lower()
        stmt = c.get("stmt") or ""
        fact = c.get("fact") or ""
        risk_text = (c.get("risk") or "").lower()

        # 1) Customer-specific CRM data → never a product contradiction
        if verdict in {"fail", "risk"} and _is_customer_specific_statement(stmt):
            c["verdict"] = "pass"
            c["vtext"] = "✓ ACCURATE"
            c["risk"] = "None"
            c["note"] = (
                "Statement is customer-specific data (e.g. their policy name, "
                "sum assured, maturity date) read back from CRM — not a "
                "product-spec claim, so not compared against the spec."
            )
            continue

        # 2) Meta-instructions in the spec are not contradicting facts
        if verdict in {"fail", "risk"} and _fact_is_meta_instruction(fact):
            c["verdict"] = "pass"
            c["vtext"] = "✓ ACCURATE"
            c["risk"] = "None"
            c["note"] = "Spec sentence was a meta-instruction (e.g. 'verify from system'), not a contradicting fact."
        elif verdict == "fail":
            has_numeric_reason = bool(re.search(r"\b(numeric|differ|conflict|contradict|wrong number|wrong age|wrong premium|wrong amount)\b", risk_text))
            has_spec_reason = bool(re.search(r"\b(spec|irda|rbi|regulator|product circular|policy bond)\b", risk_text))
            if not customer_disputed and not has_numeric_reason and not has_spec_reason:
                c["verdict"] = "risk"
                c["vtext"] = "⚠️ CONDITIONALLY OK"
                c["note"] = "Downgraded: no customer dispute and no concrete spec contradiction. Manual review recommended."

    # ------------------------------------------------------------------
    # Fail / risk signal extraction
    # ------------------------------------------------------------------
    fail_checks = [c for c in product_checks if (c.get("verdict") or "").lower() == "fail"]
    risk_checks = [c for c in product_checks if (c.get("verdict") or "").lower() == "risk"]
    hard_fail = len(fail_checks) > 0
    risk_only = (not hard_fail) and len(risk_checks) > 0

    # Confidence test: is a fail check substantiated with a concrete reason?
    def _is_substantiated(chk: Dict[str, Any]) -> bool:
        risk_text = (chk.get("risk") or "").lower()
        return bool(re.search(
            r"(numeric|conflicts?|contradict|wrong|incorrect|inaccurate|irda|rbi|specification|spec |misstate|misleading)",
            risk_text,
        ))

    substantiated_fails = [c for c in fail_checks if _is_substantiated(c)]
    model_fatal_flag = ("false_information" in flags) or ("behavior_issue" in flags)
    # ------------------------------------------------------------------
    # System-failure rule — broadened to cover "couldn't fetch / unable to
    # retrieve / failed to access" etc. ANY such bot response on a
    # product / customer-data query is a HARD system failure and the call
    # is escalated to FATAL (not just watch / behavior_issue).
    # ------------------------------------------------------------------
    product_keywords = ("policy", "product", "premium", "sum assured", "maturity",
                        "surrender", "loan", "coverage", "term", "nominee",
                        "fund value", "claim", "renewal")
    product_query = any(
        any(k in (turn.get("text", "") or "").lower() for k in product_keywords)
        for turn in turns
    )
    failed_bot_turn = next(
        (
            turn for turn in turns
            if (turn.get("speaker") or "").lower() in {"bot", "agent", "system"}
            and _is_system_failure_response(turn.get("text", "") or "")
        ),
        None,
    )
    unable_response = failed_bot_turn is not None

    # ------------------------------------------------------------------
    # Severity decision
    # ------------------------------------------------------------------
    if hard_fail:
        # Confidence gating: promote to fatal only with strong evidence.
        #   - At least one substantiated fail (spec / numeric / regulator), OR
        #   - Two independent fails, OR
        #   - Customer explicitly disputed something the bot said.
        strong_evidence = (
            model_fatal_flag and (
                len(substantiated_fails) >= 1
                or len(fail_checks) >= 2
                or customer_disputed
            )
        )

        if strong_evidence:
            analysis["severity"] = "fatal"
            analysis["fatal_reason"] = (
                analysis.get("fatal_reason")
                or ("Customer explicitly disputed bot information." if customer_disputed else None)
                or "; ".join((c.get("risk") or "").strip() for c in substantiated_fails if c.get("risk"))
                or "Substantiated product / compliance failure."
            )
            analysis["product_accuracy_score"] = 0
            analysis["pass_fail"] = "FAIL"
            if "false_information" not in flags and (
                customer_disputed or any(
                    "numeric" in (c.get("risk") or "").lower() or "spec" in (c.get("risk") or "").lower()
                    for c in substantiated_fails
                )
            ):
                flags.append("false_information")
        else:
            # Surface for human review, do NOT auto-fatal on weak signal.
            analysis["severity"] = "watch"
            analysis["product_accuracy_score"] = analysis.get("product_accuracy_score") or 2
            analysis["pass_fail"] = analysis.get("pass_fail") or "PASS"
            # Remove unsubstantiated false_information flag if the model raised
            # it without evidence (customer didn't dispute and no spec contradiction).
            if "false_information" in flags and not customer_disputed and not substantiated_fails:
                flags = [f for f in flags if f != "false_information"]

        analysis["product_issues"] = analysis.get("product_issues") or (
            "; ".join(
                f"{(c.get('stmt') or '')[:120]} → {(c.get('fact') or '')[:120]}"
                for c in fail_checks
            )
            or "Product information needs human review."
        )

        scores["response_accuracy"] = min(int(scores.get("response_accuracy", 5) or 5), 3)
        if "response_accuracy" not in failed_parameters:
            failed_parameters.append("response_accuracy")

    elif risk_only:
        current_accuracy = int(scores.get("response_accuracy", 4) or 4)
        scores["response_accuracy"] = min(current_accuracy, 3)
        if not analysis.get("product_issues") or analysis.get("product_issues") in {"None", ""}:
            analysis["product_issues"] = (
                "Unconfirmed product information should be verified against the product specification."
            )
        if analysis.get("product_accuracy_score") in {None, "", 0}:
            analysis["product_accuracy_score"] = 3
    
    # System failure rule (independent of product checks).
    # When the bot tells the customer it cannot fetch / retrieve / access the
    # requested data, the customer left the call WITHOUT what they asked for.
    # That is a hard system failure and must be FATAL, not a soft watch.
    # if unable_response and product_query:
    #     scores["system_behaviour"] = min(int(scores.get("system_behaviour", 3) or 3), 1)
    #     scores["compliance"] = min(int(scores.get("compliance", 3) or 3), 2)
    #     scores["resolution"] = min(int(scores.get("resolution", 3) or 3), 1)
    #     if "behavior_issue" not in flags:
    #         flags.append("behavior_issue")
    #     if "system_failure" not in flags:
    #         flags.append("system_failure")

    #     # Promote severity to fatal (overrides hard_fail / risk_only branch above)
    #     analysis["severity"] = "fatal"
    #     analysis["pass_fail"] = "FAIL"
    #     failure_quote = ""
    #     if failed_bot_turn:
    #         failure_quote = (failed_bot_turn.get("text") or "").strip()
    #         if len(failure_quote) > 220:
    #             failure_quote = failure_quote[:217].rstrip() + "…"
    #     analysis["fatal_reason"] = (
    #         f"System failure: bot was unable to fetch / retrieve the data the customer asked for"
    #         + (f' — "{failure_quote}"' if failure_quote else ".")
    #     )
    #     existing_issues = (analysis.get("product_issues") or "").strip()
    #     sf_issue = "System failure: bot could not fetch the requested policy data."
    #     if not existing_issues or existing_issues.lower() in {"none", ""}:
    #         analysis["product_issues"] = sf_issue
    #     elif sf_issue.lower() not in existing_issues.lower():
    #         analysis["product_issues"] = f"{existing_issues}; {sf_issue}"
    #     for fp in ("system_behaviour", "resolution"):
    #         if fp not in failed_parameters:
    #             failed_parameters.append(fp)
    if unable_response and product_query:
        escalated = _bot_escalated(turns)
        if not escalated:
            # existing fatal-promotion logic stays as-is, unchanged
            scores["system_behaviour"] = min(int(scores.get("system_behaviour", 3) or 3), 1)
            scores["compliance"] = min(int(scores.get("compliance", 3) or 3), 2)
            scores["resolution"] = min(int(scores.get("resolution", 3) or 3), 1)
            if "behavior_issue" not in flags:
                flags.append("behavior_issue")
            if "system_failure" not in flags:
                flags.append("system_failure")
            analysis["severity"] = "fatal"
            analysis["pass_fail"] = "FAIL"
            failure_quote = (failed_bot_turn.get("text") or "").strip() if failed_bot_turn else ""
            if len(failure_quote) > 220:
                failure_quote = failure_quote[:217].rstrip() + "…"
            analysis["fatal_reason"] = (
                "System failure: bot was unable to fetch / retrieve the data the customer asked for"
                + (f' — "{failure_quote}"' if failure_quote else ".")
            )
            existing_issues = (analysis.get("product_issues") or "").strip()
            sf_issue = "System failure: bot could not fetch the requested policy data."
            if not existing_issues or existing_issues.lower() in {"none", ""}:
                analysis["product_issues"] = sf_issue
            elif sf_issue.lower() not in existing_issues.lower():
                analysis["product_issues"] = f"{existing_issues}; {sf_issue}"
            for fp in ("system_behaviour", "resolution"):
                if fp not in failed_parameters:
                    failed_parameters.append(fp)
        else:
            # Correct escalation — resolution achieved via handoff, not a failure.
            # Give a small dip to system_behaviour only if NOT already a clean escalation,
            # otherwise leave scores untouched.
            analysis["category"] = analysis.get("category") or "Escalation Request"
            log.info("[QA] System-failure phrase detected but bot escalated — not penalizing.")
    # ------------------------------------------------------------------
    # Finalise
    # ------------------------------------------------------------------
    flags = list(dict.fromkeys(flags))
    analysis["scores"] = scores
    analysis["flags"] = flags
    analysis["failed_parameters"] = failed_parameters
    analysis["qa_findings"] = _build_qa_findings({**analysis, "transcript": turns})
    analysis["score_reason"] = _score_reason(analysis)

    if not analysis.get("product_accuracy_score") and product_checks:
        analysis["product_accuracy_score"] = 5 if (not risk_only and not hard_fail) else 3
    if not analysis.get("product_issues"):
        analysis["product_issues"] = "None"
    if analysis.get("severity") == "fatal" and not analysis.get("fatal_reason"):
        analysis["fatal_reason"] = "Substantiated product or policy failure."
    if analysis.get("severity") not in {"fatal", "critical", "watch", "normal"}:
        analysis["severity"] = "watch"
    if analysis.get("pass_fail") not in {"PASS", "FAIL"}:
        analysis["pass_fail"] = "FAIL" if failed_parameters else "PASS"
    # Always recompute — 4o mini's arithmetic is unreliable
    analysis["weighted_score"] = _compute_weighted_score(analysis.get("scores", {}))
    analysis["failed_parameters"] = list(dict.fromkeys(
        _compute_failed_parameters(analysis.get("scores", {})) + analysis.get("failed_parameters", [])
    ))
    analysis = _restrict_failure_criteria(analysis)
    return analysis


PRODUCT_FEATURE_KEYWORDS = {
    "premium": ["premium", "installment", "monthly", "quarterly", "half-yearly", "yearly", "single premium"],
    "benefit": ["benefit", "sum assured", "cover", "coverage", "death benefit", "maturity benefit", "payout"],
    "policy_term": ["term", "tenure", "policy term", "duration", "years"],
    "eligibility": ["eligibility", "age", "entry age", "minimum age", "maximum age"],
    "surrender": ["surrender", "paid up", "discontinuance", "withdrawal", "partial withdrawal", "lapse"],
    "claim": ["claim", "nominee", "death claim", "settlement"],
    "grace_period": ["grace period", "grace", "missed premium", "late premium"],
    "loan": ["loan", "policy loan"],
    "bonus": ["bonus", "guaranteed additions", "addition"],
    "fund_value": ["fund value", "nav", "unit", "ulip"],
    "free_look": ["free look", "free-look", "cancellation"],
    "rider": ["rider", "accidental", "critical illness"],
}


def _tokenize(text: str) -> List[str]:
    return [t for t in re.findall(r"[a-z0-9]+", (text or "").lower()) if len(t) > 2]
def _collect_products_mentioned(
    analysis: Dict[str, Any],
    rag_hit: Dict[str, Any],
    transcript_text: str = "",
) -> List[str]:
    seen: set = set()
    out: List[str] = []

    def _add(name: Any):
        if not name:
            return
        s = _safe_filename_label(str(name)).strip()
        if not s or s.lower() in {"none", "unknown"}:
            return
        key = s.lower()
        if key in seen:
            return
        seen.add(key)
        out.append(s)

    # 1. Explicit transcript names FIRST — highest confidence
    for name in _extract_product_mentions_from_text(transcript_text or ""):
        _add(name)
    for name in (rag_hit or {}).get("explicitly_mentioned", []):
        _add(name)

    # 2. GPT's answer
    _add(analysis.get("product_mentioned"))
    for p in (analysis.get("products_mentioned") or []):
        _add(p)

    # 3. RAG hits
    _add((rag_hit or {}).get("product"))
    for sec in ((rag_hit or {}).get("secondary_products") or []):
        _add(sec.get("product"))
    for name in ((rag_hit or {}).get("all_product_scores") or {}).keys():
        _add(name)

    return out

def _sentence_snippets(text: str, limit: int = 4) -> List[str]:
    sentences = re.split(r"(?<=[.!?])\s+", text or "")
    snippets = []
    for sentence in sentences:
        normalized = sentence.strip()
        if len(normalized) < 50:
            continue
        if any(keyword in normalized.lower() for keywords in PRODUCT_FEATURE_KEYWORDS.values() for keyword in keywords):
            snippets.append(normalized)
        if len(snippets) >= limit:
            break
    return snippets


def _build_product_catalog(meta_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped: Dict[str, Dict[str, Any]] = defaultdict(lambda: {"chunks": [], "sources": set()})
    for row in meta_rows:
        product = row.get("product") or "Unknown"
        grouped[product]["chunks"].append(row)
        grouped[product]["sources"].add(row.get("source", ""))

    catalog = []
    for product, data in grouped.items():
        corpus = " ".join(chunk.get("text", "") for chunk in data["chunks"])
        tokens = _tokenize(corpus)
        token_counts = Counter(tokens)
        top_terms = [term for term, _ in token_counts.most_common(24)]
        feature_hits = Counter()
        for chunk in data["chunks"]:
            chunk_text = chunk.get("text", "").lower()
            for feature, keywords in PRODUCT_FEATURE_KEYWORDS.items():
                if any(keyword in chunk_text for keyword in keywords):
                    feature_hits[feature] += 1
        snippets = _sentence_snippets(corpus, limit=5)
        catalog.append({
            "product": product,
            "sources": sorted(s for s in data["sources"] if s),
            "chunk_count": len(data["chunks"]),
            "top_terms": top_terms[:12],
            "feature_hits": dict(feature_hits.most_common(8)),
            "summary": snippets[0] if snippets else (corpus[:320] + "..." if len(corpus) > 320 else corpus),
            "evidence_snippets": snippets,
        })

    catalog.sort(key=lambda item: (item["chunk_count"], len(item["top_terms"])), reverse=True)
    return catalog


def _document_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {i}]\n{text}")
        return "\n\n".join(pages)
    if ext in {".docx", ".doc"}:
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if ext in {".txt"}:
        return path.read_text(errors="replace")
    return ""


def _chunk_text(text: str, chunk_size: int = 400, overlap: int = 100,
                sentence_mode: bool = True) -> List[str]:
    """Sentence-aware chunking optimised for product specs.

    When `sentence_mode` is True (default) we emit ONE chunk per meaningful
    sentence (≥ 12 words). Adjacent short sentences are merged until the
    combined chunk reaches a minimum length. This keeps numeric facts intact
    (premium amounts, ages, %), which previously caused false 'numeric
    conflict' fails when facts straddled chunk boundaries.

    A single huge sentence (> chunk_size words) is split with word-level
    overlap so it is not lost.

    Args:
        chunk_size: maximum word count per chunk.
        overlap:    word count carried over between consecutive chunks.
        sentence_mode: if False, falls back to the older grouped behaviour.
    """
    text = (text or "").strip()
    if not text:
        return []

    # Split on sentence boundaries while preserving punctuation.
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)
    sentences = [s.strip() for s in sentences if s.strip()]
    if not sentences:
        return []

    MIN_WORDS = 12          # below this, sentence is "short" and gets merged
    MIN_CHUNK_CHARS = 80    # discard tiny residuals
    chunks: List[str] = []

    if sentence_mode:
        buffer: List[str] = []
        buffer_len = 0
        for sent in sentences:
            sent_words = sent.split()
            sent_len = len(sent_words)

            # Huge sentence — split with word overlap
            if sent_len > chunk_size:
                if buffer:
                    chunks.append(" ".join(buffer).strip())
                    buffer, buffer_len = [], 0
                step = max(1, chunk_size - overlap)
                for start in range(0, sent_len, step):
                    piece = " ".join(sent_words[start:start + chunk_size]).strip()
                    if len(piece) >= MIN_CHUNK_CHARS:
                        chunks.append(piece)
                continue

            # Substantial standalone sentence — flush buffer + emit as its own chunk
            if sent_len >= MIN_WORDS and buffer_len == 0:
                chunks.append(sent.strip())
                continue

            # Short sentence — accumulate
            if buffer_len + sent_len <= chunk_size:
                buffer.append(sent)
                buffer_len += sent_len
            else:
                if buffer:
                    chunks.append(" ".join(buffer).strip())
                buffer = [sent]
                buffer_len = sent_len

            # If the accumulated buffer is now substantial, flush it.
            if buffer_len >= MIN_WORDS:
                chunks.append(" ".join(buffer).strip())
                buffer, buffer_len = [], 0

        if buffer:
            tail = " ".join(buffer).strip()
            if len(tail) >= MIN_CHUNK_CHARS or not chunks:
                chunks.append(tail)
        return [c for c in chunks if len(c) >= MIN_CHUNK_CHARS]

    # Legacy grouped mode (fallback)
    current: List[str] = []
    current_len = 0
    for sent in sentences:
        sent_words = sent.split()
        sent_len = len(sent_words)
        if sent_len > chunk_size:
            if current:
                chunks.append(" ".join(current).strip())
                current, current_len = [], 0
            step = max(1, chunk_size - overlap)
            for start in range(0, sent_len, step):
                piece = " ".join(sent_words[start:start + chunk_size]).strip()
                if len(piece) >= MIN_CHUNK_CHARS:
                    chunks.append(piece)
            continue
        if current_len + sent_len <= chunk_size:
            current.append(sent)
            current_len += sent_len
        else:
            if current:
                chunks.append(" ".join(current).strip())
            if overlap > 0 and current:
                tail_words = " ".join(current).split()[-overlap:]
                current = [" ".join(tail_words), sent]
                current_len = len(tail_words) + sent_len
            else:
                current = [sent]
                current_len = sent_len
    if current:
        tail = " ".join(current).strip()
        if len(tail) >= MIN_CHUNK_CHARS or not chunks:
            chunks.append(tail)
    return [c for c in chunks if len(c) >= MIN_CHUNK_CHARS]


def _collect_product_files() -> List[Path]:
    files = []
    seen = set()
    for source_dir in PRODUCT_SOURCES:
        if not source_dir.exists():
            continue
        for path in sorted(source_dir.glob("*")):
            if path.suffix.lower() not in {".pdf", ".docx", ".doc", ".txt"}:
                continue
            try:
                digest = hashlib.sha1(path.read_bytes()).hexdigest()
            except Exception:
                digest = f"{path.name}:{path.stat().st_size if path.exists() else 0}"
            if digest in seen:
                continue
            seen.add(digest)
            files.append(path)
    return files


def _save_rag_backend(mode: str):
    RAG_BACKEND_FILE.write_text(json.dumps({"mode": mode, "updated_at": datetime.now().isoformat()}, indent=2))


def _load_rag_meta() -> List[Dict[str, Any]]:
    try:
        return json.loads(RAG_META_FILE.read_text())
    except Exception:
        return []


def _load_rag_index():
    if faiss is None or not RAG_INDEX_FILE.exists() or not RAG_META_FILE.exists():
        return None, []
    try:
        index = faiss.read_index(str(RAG_INDEX_FILE))
        meta = _load_rag_meta()
        return index, meta
    except Exception as e:
        log.warning(f"FAISS index load failed: {e}")
        return None, []

def rebuild_product_rag_index() -> dict:
    files = _collect_product_files()
    meta_rows: List[Dict[str, Any]] = []

    if not files:
        if RAG_INDEX_FILE.exists():
            RAG_INDEX_FILE.unlink(missing_ok=True)
        RAG_META_FILE.write_text("[]")
        _save_rag_backend("empty")
        sig_path = PROC_DIR / "rag_signature.json"
        if sig_path.exists():
            sig_path.unlink(missing_ok=True)
        return {"mode": "empty", "chunks": 0, "products": 0}

    sig_path = PROC_DIR / "rag_signature.json"
    embed_cache_path = PROC_DIR / "embed_cache.npz"
    current_sig: Dict[str, str] = {}
    for path in files:
        try:
            current_sig[path.name] = hashlib.sha1(path.read_bytes()).hexdigest()
        except Exception:
            current_sig[path.name] = f"{path.name}:{path.stat().st_size}"

    # ---- (signature short-circuit removed — always rebuild meta_rows) ----

    # ---- load existing embedding cache (key: sha1(chunk text)) ----
    embed_cache: Dict[str, np.ndarray] = {}
    if embed_cache_path.exists():
        try:
            with np.load(str(embed_cache_path), allow_pickle=False) as npz:
                cache_keys = npz["keys"].tolist() if "keys" in npz else []
                cache_mat = npz["vecs"] if "vecs" in npz else np.zeros((0, 0))
                for i, k in enumerate(cache_keys):
                    embed_cache[k] = cache_mat[i]
            log.info(f"Loaded {len(embed_cache)} cached embeddings.")
        except Exception as e:
            log.warning(f"Embed cache load failed (rebuilding): {e}")
            embed_cache = {}

    try:
        model = _load_embedder()
        chunk_texts: List[str] = []
        chunk_keys: List[str] = []
        for path in files:
            doc_text = _document_text(path)
            if not doc_text.strip():
                continue
            chunks = _chunk_text(doc_text)
            product_label = _safe_filename_label(path.name)
            for chunk_idx, chunk in enumerate(chunks):
                chunk_texts.append(chunk)
                chunk_keys.append(hashlib.sha1(chunk.encode("utf-8")).hexdigest())
                meta_rows.append({
                    "source": path.name,
                    "path": str(path),
                    "product": product_label,
                    "chunk_index": chunk_idx,
                    "text": chunk,
                    "word_count": len(chunk.split()),
                    "page": None,
                })

        if not chunk_texts:
            RAG_META_FILE.write_text("[]")
            _save_rag_backend("empty")
            sig_path.write_text(json.dumps(current_sig, indent=2))
            return {"mode": "empty", "chunks": 0, "products": len(files)}

        missing_idx = [i for i, k in enumerate(chunk_keys) if k not in embed_cache]
        if missing_idx:
            log.info(f"Embedding {len(missing_idx)} new chunks (cached: {len(chunk_keys) - len(missing_idx)}).")
            new_vecs = _embed_texts([chunk_texts[i] for i in missing_idx])
            for vec, i in zip(new_vecs, missing_idx):
                embed_cache[chunk_keys[i]] = vec.astype(np.float32)
        else:
            log.info("All chunks already cached — no embedding calls made.")

        vectors = np.vstack([embed_cache[k] for k in chunk_keys]).astype(np.float32)

        if faiss is not None:
            index = faiss.IndexFlatIP(vectors.shape[1])
            index.add(vectors)
            faiss.write_index(index, str(RAG_INDEX_FILE))
            _save_rag_backend("faiss")
        else:
            RAG_INDEX_FILE.write_text("")
            _save_rag_backend("keyword")

        RAG_META_FILE.write_text(json.dumps(meta_rows, indent=2, default=str))

        sig_path.write_text(json.dumps(current_sig, indent=2))
        try:
            keys_arr = np.array(list(embed_cache.keys()))
            vecs_arr = np.vstack(list(embed_cache.values())).astype(np.float32)
            np.savez_compressed(str(embed_cache_path), keys=keys_arr, vecs=vecs_arr)
        except Exception as e:
            log.warning(f"Embed cache write failed: {e}")

        return {
            "mode": "faiss" if faiss is not None else "keyword",
            "chunks": len(meta_rows),
            "products": len(files),
            "cache": "partial" if missing_idx else "full",
            "embedded_now": len(missing_idx),
        }
    except Exception as e:
        log.warning(f"Product RAG rebuild failed: {e}")
        RAG_META_FILE.write_text("[]")
        _save_rag_backend("keyword")
        return {"mode": "keyword", "chunks": 0, "products": len(files), "error": str(e)}
        
def _keyword_rank(query: str, meta_rows: List[Dict[str, Any]], top_k: int = 5) -> List[Dict[str, Any]]:
    terms = [t for t in re.findall(r"[a-z0-9]+", query.lower()) if len(t) > 2]
    if not terms:
        return meta_rows[:top_k]

    scored = []
    for idx, row in enumerate(meta_rows):
        text = row.get("text", "").lower()
        score = sum(2 if term in row.get("product", "").lower() else 1 for term in terms if term in text or term in row.get("product", "").lower())
        if score:
            scored.append((score, idx, row))
    scored.sort(key=lambda x: (x[0], x[1]), reverse=True)
    return [row for _, _, row in scored[:top_k]]


def search_product_rag(query: str, top_k: int = 5) -> List[Dict[str, Any]]:
    query = (query or "").strip()
    if not query:
        return []

    # Check cache first
    cache_key = f"rag_search:{query}:{top_k}"
    cached = get_cache(cache_key)
    if cached is not None:
        return cached

    index, meta_rows = _load_rag_index()
    if faiss is not None and index is not None and meta_rows:
        try:
            query_vec = _embed_texts([query])
            scores, idxs = index.search(query_vec, min(top_k, len(meta_rows)))
            rows = []
            for score, idx in zip(scores[0].tolist(), idxs[0].tolist()):
                if idx < 0 or idx >= len(meta_rows):
                    continue
                row = dict(meta_rows[idx])
                row["score"] = float(score)
                rows.append(row)
            if rows:
                set_cache(cache_key, rows)
                return rows
        except Exception as e:
            log.warning(f"FAISS retrieval failed, using keyword fallback: {e}")

    meta_rows = _load_rag_meta()
    if not meta_rows:
        set_cache(cache_key, [])
        return []
    ranked = _keyword_rank(query, meta_rows, top_k=top_k)
    for i, row in enumerate(ranked):
        row["score"] = float(max(0.0, 1.0 - i * 0.1))
    set_cache(cache_key, ranked)
    return ranked

def infer_product_context(transcript_text: str, top_k: int = 5) -> Dict[str, Any]:
    query = re.sub(r"\s+", " ", transcript_text or "").strip()
    if not query:
        return {"product": "None", "context": "", "chunks": []}

    cache_key = f"product_context:{hashlib.md5(query.encode()).hexdigest()}"
    cached = get_cache(cache_key)
    if cached is not None:
        return cached

    rows = search_product_rag(query, top_k=max(top_k, 12))
    meta_rows = _load_rag_meta()
    catalog = _build_product_catalog(meta_rows) if meta_rows else []

    if not rows and not catalog:
        result = {"product": "None", "context": "", "chunks": []}
        set_cache(cache_key, result)
        return result

    transcript_lower = query.lower()
    transcript_tokens = set(_tokenize(query))
    product_scores: Dict[str, float] = defaultdict(float)
    product_evidence: Dict[str, list] = defaultdict(list)
    product_matches: Dict[str, set] = defaultdict(set)

    # Step 1: RAG similarity scores (unchanged)
    for rank, row in enumerate(rows):
        product = row.get("product", "Unknown")
        score = float(row.get("score", 0.0))
        product_scores[product] += max(score, 0.0) + (0.05 * (len(rows) - rank))
        product_evidence[product].append(row)

    for profile in catalog:
        product = profile["product"]
        top_terms = set(profile.get("top_terms", []))
        overlap = transcript_tokens.intersection(top_terms)
        if overlap:
            product_scores[product] += 0.12 * len(overlap)
            product_matches[product].update(sorted(overlap))
        for feature, keywords in PRODUCT_FEATURE_KEYWORDS.items():
            if any(keyword in query.lower() for keyword in keywords):
                if feature in profile.get("feature_hits", {}):
                    product_scores[product] += 0.35 + 0.05 * float(
                        profile.get("feature_hits", {}).get(feature, 0))
                    product_matches[product].add(feature)

    # ── Step 2: EXPLICIT NAME BOOST ──────────────────────────────────────────
    EXPLICIT_MENTION_BONUS = 50.0  # always beats RAG similarity

    def _normalise(s: str) -> str:
        s = re.sub(r"(?i)\bbajaj\s+allianz\s+life\b", "bajaj life", s)
        return re.sub(r"\s+", " ", s.lower()).strip()

    name_to_product: Dict[str, str] = {}
    for profile in catalog:
        prod = profile["product"]
        norm = _normalise(prod)
        name_to_product[norm] = prod
        # Short alias: "eTouch2" without the brand prefix
        short = re.sub(r"^bajaj\s+(?:allianz\s+)?(?:life\s+)?", "", norm).strip()
        if short and short != norm:
            name_to_product[short] = prod

    explicitly_mentioned: set = set()

    for norm_name, prod in name_to_product.items():
        if norm_name and norm_name in transcript_lower:
            product_scores[prod] += EXPLICIT_MENTION_BONUS
            product_matches[prod].add("explicit_name_mention")
            explicitly_mentioned.add(prod)
            log.info(f"[RAG] Explicit mention boost: '{prod}' (+{EXPLICIT_MENTION_BONUS})")

    # Also catch names found by the transcript regex extractor
    for tm in _extract_product_mentions_from_text(transcript_text or ""):
        tm_norm = _normalise(tm)
        matched = next(
            (prod for norm, prod in name_to_product.items()
             if norm in tm_norm or tm_norm in norm),
            None,
        )
        if matched:
            product_scores[matched] += EXPLICIT_MENTION_BONUS
            explicitly_mentioned.add(matched)

    if not product_scores:
        result = {"product": "None", "context": "", "chunks": []}
        set_cache(cache_key, result)
        return result

    # Step 3: rank
    best_product, best_score = max(product_scores.items(), key=lambda x: x[1])
    best_profile = next((p for p in catalog if p["product"] == best_product), {})
    supporting_rows = sorted(
        product_evidence.get(best_product, []),
        key=lambda r: r.get("score", 0.0), reverse=True
    )[:top_k]

    context_lines = []
    for row in supporting_rows:
        context_lines.append(
            f"[{row.get('product','Unknown')} | {row.get('source','spec')} "
            f"| chunk {row.get('chunk_index',0)}] {row.get('text','')}"
        )
    if best_profile:
        context_lines.append(
            f"[PROFILE | {best_profile.get('product','Unknown')}] "
            f"Summary: {best_profile.get('summary','')}\n"
            f"Feature anchors: "
            f"{', '.join(sorted(best_profile.get('feature_hits',{}).keys())[:8]) or 'None'}"
        )

    # Step 4: multi-product — explicitly mentioned products are ALWAYS included
    secondary_products: List[Dict[str, Any]] = []
    for prod, score in product_scores.items():
        if prod == best_product:
            continue
        is_explicit = prod in explicitly_mentioned
        is_close    = score >= best_score * 0.55
        if not (is_explicit or is_close):
            continue
        sec_profile = next((p for p in catalog if p["product"] == prod), {})
        sec_rows = sorted(
            product_evidence.get(prod, []),
            key=lambda r: r.get("score", 0.0), reverse=True
        )[:max(2, top_k // 2)]
        secondary_products.append({
            "product": prod,
            "score": round(float(score), 3),
            "explicit_mention": is_explicit,
            "evidence": sec_rows,
            "profile": sec_profile,
        })
        for row in sec_rows:
            context_lines.append(
                f"[SECONDARY | {row.get('product','Unknown')} | "
                f"{row.get('source','spec')} | chunk {row.get('chunk_index',0)}] "
                f"{row.get('text','')}"
            )

    secondary_products.sort(key=lambda x: (-int(x["explicit_mention"]), -x["score"]))
    secondary_products = secondary_products[:5]

    evidence_terms = sorted(product_matches.get(best_product, set()))[:12]
    confidence = round(float(best_score), 3)
    product_confidence = min(0.99, 0.4 + confidence / (EXPLICIT_MENTION_BONUS + 4.0))
    if best_product in explicitly_mentioned:
        product_confidence = min(0.99, product_confidence + 0.45)

    result = {
        "product": best_product if best_score > 0.08 else "None",
        "confidence": round(product_confidence, 3),
        "matched_terms": evidence_terms,
        "explicitly_mentioned": list(explicitly_mentioned),
        "catalog": catalog[:8],
        "evidence": supporting_rows,
        "context": "\n\n".join(context_lines),
        "chunks": supporting_rows,
        "profile": best_profile,
        "secondary_products": secondary_products,
        "all_product_scores": {
            p: round(float(s), 3)
            for p, s in sorted(product_scores.items(), key=lambda x: -x[1])[:8]
        },
    }
    set_cache(cache_key, result)
    return result

def query_product_rag(product_name: str, question: str) -> str:
    query = " ".join(part for part in [product_name, question] if part).strip()
    if not query:
        return "No product specifications loaded."
    rows = search_product_rag(query, top_k=5)
    if not rows:
        return "No matching product spec found."
    return "\n\n".join(
        f"[{row.get('product', 'Unknown')} | {row.get('source', 'spec')}] {row.get('text', '')}"
        for row in rows[:3]
    )


def remove_product_from_index(filename: str):
    safe_name = Path(filename).name
    target = PRODUCT_DIR / safe_name
    if target.exists():
        target.unlink()
    rebuild_product_rag_index()


def remove_all_product_specs():
    for path in PRODUCT_DIR.glob("*"):
        if path.suffix.lower() in {".pdf", ".docx", ".doc", ".txt"}:
            path.unlink(missing_ok=True)
    rebuild_product_rag_index()

# ── Transcript Parsers ───────────────────────────────────────────────────────

def parse_transcript_text(raw: str) -> List[Dict]:
    """Parse raw transcript text into list of {sl, speaker, text}"""
    turns = []
    # Pattern: optional SL number, speaker label, colon, then quoted or unquoted text
    pattern = re.compile(
        r'(?:(\d+)\.\s+)?'           # optional SL no
        r'(bot|agent|assistant|customer|user|ivr|system)[:\s]+'  # speaker
        r'["\u201c]?(.+?)["\u201d]?(?=(?:\d+\.)?\s*(?:bot|agent|customer|user|ivr|system)[:\s]|$)',
        re.IGNORECASE | re.DOTALL
    )
    for m in pattern.finditer(raw):
        sl, speaker, text = m.group(1), m.group(2).strip(), m.group(3).strip()
        turns.append({
            "sl": int(sl) if sl else len(turns)+1,
            "speaker": speaker.lower(),
            "text": text.replace('\n', ' ').strip()
        })
    if not turns:
        # fallback: split by newlines, detect speaker
        for i, line in enumerate(raw.split('\n')):
            line = line.strip()
            if not line:
                continue
            m2 = re.match(r'^(?:\d+\.\s+)?(bot|agent|customer|user)[:\s]+["\u201c]?(.+)["\u201d]?$', line, re.I)
            if m2:
                turns.append({"sl": i+1, "speaker": m2.group(1).lower(), "text": m2.group(2).strip()})
    return turns

def extract_from_pdf(path: Path) -> List[Dict[str, str]]:
    """Returns list of {name, text} dicts (one per transcript in the PDF)"""
    reader = PdfReader(str(path))
    full_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    # Try to split by call markers
    splits = re.split(r'(?:Call\s*#?\d+|Transcript\s*\d+|={5,}|-{5,})', full_text, flags=re.IGNORECASE)
    results = []
    for i, chunk in enumerate(splits):
        chunk = chunk.strip()
        if len(chunk) > 100:
            results.append({"name": f"{path.stem}_call_{i+1}", "text": chunk})
    if not results:
        results = [{"name": path.stem, "text": full_text}]
    return results

def extract_from_docx(path: Path) -> List[Dict[str, str]]:
    doc = DocxDocument(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    splits = re.split(r'(?:Call\s*#?\d+|Transcript\s*\d+)', full_text, flags=re.IGNORECASE)
    results = []
    for i, chunk in enumerate(splits):
        chunk = chunk.strip()
        if len(chunk) > 100:
            results.append({"name": f"{path.stem}_call_{i+1}", "text": chunk})
    if not results:
        results = [{"name": path.stem, "text": full_text}]
    return results

def extract_from_excel(path: Path) -> List[Dict[str, str]]:
    """Excel may have SL column with multiple transcripts"""
    results = []
    try:
        xl = pd.ExcelFile(str(path))
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            df.columns = [str(c).strip().lower() for c in df.columns]
            # Detect SL column
            sl_col = next((c for c in df.columns if re.match(r'sl|s\.no|serial|#', c)), None)
            # Detect transcript column
            tx_col = next((c for c in df.columns if any(k in c for k in ['transcript','text','conversation','call'])), None)
            if sl_col and tx_col:
                for _, row in df.dropna(subset=[tx_col]).iterrows():
                    sl = str(row.get(sl_col, '')).strip()
                    text = str(row[tx_col]).strip()
                    if len(text) > 50:
                        name = f"{path.stem}_sl{sl}" if sl else f"{path.stem}_row{_}"
                        results.append({"name": name, "text": text, "sl": sl})
            elif tx_col:
                for i, row in df.dropna(subset=[tx_col]).iterrows():
                    text = str(row[tx_col]).strip()
                    if len(text) > 50:
                        results.append({"name": f"{path.stem}_row{i+1}", "text": text})
            else:
                # Treat each row as part of one transcript
                all_text = "\n".join(str(v) for v in df.values.flatten() if pd.notna(v) and str(v).strip())
                if len(all_text) > 50:
                    results.append({"name": f"{path.stem}_{sheet}", "text": all_text})
    except Exception as e:
        log.error(f"Excel parse error: {e}")
    return results

def extract_transcripts_from_file(path: Path) -> List[Dict[str, str]]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_from_pdf(path)
    elif ext in [".docx", ".doc"]:
        return extract_from_docx(path)
    elif ext in [".xlsx", ".xls"]:
        return extract_from_excel(path)
    elif ext == ".txt":
        text = path.read_text(errors='replace')
        return [{"name": path.stem, "text": text}]
    elif ext == ".json":
        return extract_from_json(path)
    return []

def extract_from_json(path: Path) -> List[Dict[str, str]]:
    """Parse a JSON file containing one or many call payloads in the ingest format."""
    role_map = {"assistant": "bot", "user": "customer"}
    try:
        raw = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception as e:
        log.error(f"JSON parse error {path}: {e}")
        return []

    # Accept either a single dict or an array of dicts
    records = raw if isinstance(raw, list) else [raw]
    items = []
    for rec in records:
        if not isinstance(rec, dict):
            continue
        # Normalise: support both unique_call_id (vendor) and sender_id (extracted_logs)
        call_id = (
            rec.get("unique_call_id")
            or rec.get("sender_id")
            or str(uuid.uuid4())
        )
        conv_log = rec.get("conversation_log", [])
        if not conv_log:
            continue
        turns = []
        for i, msg in enumerate(conv_log, start=1):
            if not isinstance(msg, dict):
                continue
            role = msg.get("role", "user")
            content = (msg.get("content") or "").strip()
            if not content:
                continue
            turns.append({
                "sl": i,
                "speaker": role_map.get(role, "unknown"),
                "text": content,
            })
        if not turns:
            continue
        raw_text = "\n".join(f"{t['speaker']}: {t['text']}" for t in turns)
        items.append({
            "name": str(call_id),
            "sl": str(call_id),
            "text": raw_text,
            "source_file": path.name,
            "turns": turns,
            "meta": {
                "intent": rec.get("INTENT"),
                "stage_code": rec.get("STAGE_CODE"),
                "sentiment": rec.get("sentiment"),
                "summary": rec.get("conversation_summary"),
                "no_of_queries": rec.get("no_of_queries"),
                "no_of_queries_resolved": rec.get("no_of_queries_resolved"),
                "call_comp_flag": rec.get("call_comp_flag"),
                "customer_journey": rec.get("customer_journey", []),
                "standalone_type": rec.get("standalone_type", []),
                "sub_query_type": rec.get("sub_query_type", []),
                "root_dscn": rec.get("root_dscn"),
                "mid_dscn": rec.get("mid_dscn"),
                "root_tta": rec.get("root_tta"),
                "mid_tta": rec.get("mid_tta"),
                "audio_link": rec.get("audio_link") or "",
            },
        })
    log.info(f"JSON ingest: {path.name} → {len(items)} calls extracted")
    return items

# ── RAG: Product Spec Retrieval ──────────────────────────────────────────────
# The local FAISS-backed helpers above handle retrieval and indexing.

# ── GPT-4o Analysis ──────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert quality assurance analyst for Bajaj Life Insurance's AI-powered customer service bot. 
Your role is to evaluate call transcripts between the BOT and customers with extreme precision and professional rigor.
IMPORTANT — you are running as a smaller, faster model. Be literal and conservative:
- If the bot escalates, transfers, or offers a callback after saying it cannot do something, this is a SUCCESSFUL resolution via escalation — do NOT flag behavior_issue or system_failure, and do NOT lower system_behaviour/resolution scores for this reason.
- Only use flags ["false_information"] if you can quote the exact contradicting product-spec sentence.
- Only use flags ["behavior_issue"] or "system_failure" if the bot failed to provide requested information AND did not escalate/transfer/offer callback.
- If a product name is stated verbatim anywhere in the transcript, you MUST set product_mentioned to that exact name — do not substitute a different product even if spec context seems to favor another.
- pass_fail should be "FAIL" only if: system_failure occurred with no escalation, OR false_information occurred, OR communication_quality score is below 3 (rude/unclear). Low scores on other parameters alone do not justify FAIL.
You will evaluate each call on 10 parameters and provide JSON output ONLY — no prose outside the JSON.

You must think like a strict insurance QA lead, not a generic chatbot. Use the transcript, customer intent, and product knowledge context together to infer which product is being discussed even when the product name is not spoken explicitly. Product identification must rely on product details, feature language, policy behavior, and benefit structure.

EVALUATION PARAMETERS (score each 1-5):
1. greeting_opening (Weight: 5%, Min pass: 3) - Polite greeting, professional tone, respectful language
2. query_understanding (Weight: 10%, Min pass: 3) - Correct interpretation of customer intent, clarification when needed
3. response_accuracy (Weight: 25%, Min pass: 4) - Factual correctness, no guessing, no contradictions, reliable info
4. communication_quality (Weight: 8%, Min pass: 3) - Clear sentences, well-organized, avoids jargon, structured
5. compliance (Weight: 20%, Min pass: 4) - Follows regulatory rules (insurance/finance/privacy), no restricted info shared, no outcome promises
6. personalisation (Weight: 5%, Min pass: 3) - Uses available info (name, SR no, context) within compliance limits
7. empathy_soft_skills (Weight: 5%, Min pass: 3) - Acknowledges feelings, warm language, avoids robotic responses
8. resolution (Weight: 10%, Min pass: 3) - Directly solves problem, clear guidance, alternatives offered
9. system_behaviour (Weight: 10%, Min pass: 3) - No loops, no latency issues, no errors, appropriate escalation
10. closing_interaction (Weight: 2%, Min pass: 3) - Polite closure, thanks customer, offers further help

PRODUCT EVALUATION:
- Identify the insurance product(s) mentioned or strongly implied by the call, even if the call never states the product name
- Analyze the actual conversation for product feature keywords: premium payment terms, coverage/sum assured amounts, eligibility criteria (age, income, employment status), policy tenure/term length, maturity/death benefits, riders, surrender/loan options, grace periods
- Use product details such as premium structure, term, eligibility, surrender/loan behavior, maturity/death benefit, grace period, riders, and payout style to infer the right product
- Explain why the chosen product matches the call and which product details support the match
- Check if product details (premium, coverage, terms, exclusions, eligibility) mentioned are accurate based on the knowledge base
- Note any product information gaps, invented facts, or contradictions
- Note what should have been said differently if the bot missed a product-specific rule or detail
- Include 2-4 product checks with exact bot statement, supporting fact from the knowledge base, and verdict
- If the bot gives incorrect product or policy information, treat it as a fatal QA defect and score product accuracy 0
- If the bot states unconfirmed product information, mark it as a risk and lower confidence instead of passing it
- If the bot says it is unable to provide information for a policy/product question, flag it as a system failure and lower the relevant QA scores
- Every product check must compare the call statement against an exact sentence from the product PDF or indexed product spec

PRODUCT CHECK VERDICT RULES — READ CAREFULLY (these matter for false-positive control):
- "pass"  → statement is consistent with the spec OR there is no spec sentence available to contradict it. When the bot quotes a SYSTEM-RETRIEVED value (e.g. "your premium is six thousand monthly", "your policy ends on DD/MM/YYYY", "your nominee is X"), default to "pass" UNLESS the spec contradicts it with an explicit different value OR the customer disputes it in the same call.
- "risk"  → bot's statement is unconfirmed by spec but plausible (e.g. value falls inside a spec range, spec is silent, bot used categorical language like "always/never/guaranteed").
- "fail"  → bot's statement DIRECTLY CONTRADICTS a specific value in the spec (e.g. spec says "entry age 18-65" and bot says "entry age is 70"), OR the customer explicitly said the bot was wrong, OR the bot promised an unauthorised outcome.
- DO NOT label as "fail" because the spec says things like "verify from system" / "refer to policy bond" / "as per records" — those are meta-instructions, not contradicting facts.
- DO NOT label as "fail" because the spec is silent on the topic — that is "risk" at most.
- DO NOT label as "fail" because the bot quoted a number with no matching number in the spec — the bot may have read from the customer's policy in CRM.
CRITICAL MULTI-PRODUCT AND SPEC-MATCHING RULES:
- A single call can involve MORE THAN ONE product (e.g. eTouch AND Goal Suraksha).
  Identify ALL products explicitly named in the transcript.
- For EACH identified product, run separate product checks using ONLY that product's spec.
  NEVER check a statement about Product A against the spec of Product B.
- Set the "product" field on every product_check entry to identify which product it belongs to.
- If the transcript explicitly states a product name (e.g. "Bajaj Life eTouch2"), that IS the
  product_mentioned regardless of what the RAG context scores suggest.
- If no spec is available in the context for a mentioned product, omit checks for it rather
  than substituting another product's spec.
- The "product_mentioned" field must reflect what the customer/bot explicitly said,
  not what the RAG context ranked highest.

MULTI-PRODUCT HANDLING:
- If the transcript mentions or strongly implies MORE THAN ONE product, identify ALL of them.
- Put the most-discussed product in "product_mentioned" and the full ordered list in "products_mentioned".
- When generating product_checks, attribute each check to the right product by referencing the product name in the "fact" or by adding a "product" field to the check.

CALL CLASSIFICATION:
- category: One of [Premium Receipt, Premium Payment Assistance, Policy Status Inquiry, Policy Document Request, Surrender Request, Loan Against Policy, Nominee Update, Address or Contact Update, Maturity Claim, Death Claim, Revival of Lapsed Policy, Free Look Cancellation, Rider Inquiry, Fund Switch or Redirection, Partial Withdrawal, Benefit Illustration Request, Complaint or Grievance, Escalation Request, General Inquiry]
- severity: "normal", "watch", "critical", "fatal"
- fatal_reason: explain if fatal (empty string otherwise)
- flags: array of ONLY actual violations: ["compliance_breach", "false_information", "regulatory_violation"] for critical issues. DO NOT flag routine statements, customer observations, or normal bot responses. ONLY flag if the bot actually violated compliance rules, gave factually incorrect product info, or exhibited problematic behavior
- sentiment: overall customer sentiment: "positive", "neutral", "frustrated", "angry", "distressed"

FLAGGING RULES - BE STRICT AND PRECISE:
- "compliance_breach" ONLY if: privacy laws violated, promised guaranteed outcomes not authorized, shared restricted financial data, violated consumer protection act
- "false_information" ONLY if BOTH of these are true:
    (a) the bot stated a product/policy fact that directly contradicts the product spec (wrong premium amount, wrong age limit, wrong benefit amount, etc) OR the customer explicitly said the bot was wrong in the same call ("that's incorrect", "no that's wrong"); AND
    (b) you can quote the contradicting spec sentence OR the customer dispute turn.
  NEVER raise false_information just because the bot quoted a number/date/name that you could not verify against the spec — the bot likely retrieved it from the CRM/system.
- "regulatory_violation" ONLY if: violated RBI/IRDAI rules, illegal solicitation, improper disclosure
- "behavior_issue" ONLY if: the bot says it CANNOT access the data the customer asked for (system failure), OR escalation needed but call ended without escalating, OR customer clearly in distress with no empathy shown, OR obvious system loop / malfunction
- DO NOT flag for: customer saying "I already have a policy", "I see one product purchased", "I'm looking for", normal questions, normal observations
- DO NOT flag for: vague responses unless they violate compliance or are about product facts
- DO NOT flag for: missing information unless the bot explicitly said something false or violated a rule
- Empty flags array is correct if none of the above violations occurred

FATAL TRIGGERS (mark fatal + flag immediately):
- Any regulatory/compliance violation (flags: ["regulatory_violation", "compliance_breach"])
- False or misleading product information given to customer (flags: ["false_information"])
- Customer personal/financial data mishandled (flags: ["compliance_breach"])
- Agent (bot) promises outcomes not authorized (flags: ["compliance_breach"])
- Customer clearly distressed with no empathy or escalation (flags: ["behavior_issue"])
- Inappropriate or offensive language (flags: ["behavior_issue"])
- Call ends without resolution on critical issue with no escalation (flags: ["behavior_issue"])

OUTPUT FORMAT — respond with ONLY this JSON structure:
{
  "scores": {
    "greeting_opening": <1-5>,
    "query_understanding": <1-5>,
    "response_accuracy": <1-5>,
    "communication_quality": <1-5>,
    "compliance": <1-5>,
    "personalisation": <1-5>,
    "empathy_soft_skills": <1-5>,
    "resolution": <1-5>,
    "system_behaviour": <1-5>,
    "closing_interaction": <1-5>
  },
    "param_comments": ["exactly 10 one-line explanations (same order as scores), each stating why that score was assigned with transcript-grounded evidence"],
  "weighted_score": <0-100 float>,
  "pass_fail": <"PASS" or "FAIL">,
  "failed_parameters": [<list of parameter names that didn't meet minimum>],
  "category": "<category>",
  "severity": "<normal|watch|critical|fatal>",
  "fatal_reason": "<string or empty string>",
  "flags": [<array of flag strings: only actual compliance_breach, false_information, regulatory_violation, or behavior_issue if they occurred; empty array if no violations>],
  "sentiment": "<positive|neutral|frustrated|angry|distressed>",
  "product_mentioned": "<primary product name or 'None'>",
  "products_mentioned": [<ALL products explicitly named in the transcript, in order of first mention>; empty array if none>],
    "product_confidence": <0-1 float>,
    "product_signals": ["keywords or product details that drove the match"],
  "product_accuracy_score": <1-5 or null>,
  "product_issues": "<string describing inaccuracies or 'None'>",
  "what_should_have_been_said": "<specific suggestions for improvement>",
  "strengths": "<what the bot did well>",
  "summary": "<2-3 sentence overall call summary>",
    "product_checks": [{"call":"<call id if available>","product":"<which product this check refers to>","stmt":"<exact bot statement>","fact":"<knowledge base fact>","verdict":"pass|fail|risk","vtext":"✓ ACCURATE|✗ INACCURATE|⚠️ CONDITIONALLY OK","risk":"None|🚨 HIGH — reason|⚠️ Medium — reason"}],
  "turn_count": <integer>,
  "bot_turns": <integer>,
  "customer_turns": <integer>,
  "estimated_duration_minutes": <float>
}"""

async def analyze_call_with_gpt4o(transcript_text: str, turns: List[Dict], product_context: str = "") -> Dict:
    """Send transcript to GPT-4o for analysis"""
    
    # Format turns for the prompt
    formatted = "\n".join(
        f"{t.get('sl', i+1)}. {t['speaker'].upper()}:\n\"{t['text']}\""
        for i, t in enumerate(turns)
    ) if turns else transcript_text
    
    product_section = f"\n\nPRODUCT SPEC CONTEXT (from RAG):\n{product_context}" if product_context else ""
    
    user_msg = f"""Analyze this Bajaj Life Insurance bot call transcript:

{formatted}
{product_section}

PRODUCT DETECTION INSTRUCTIONS:
If product specs are provided above, use them to identify which product is being discussed.
Match the conversation against the product features, benefits, policy terms, and coverage details.
Look for keywords like: premium payment mode, coverage amount, policy term, age eligibility, maturity/death benefits, riders, surrender options.
If the RAG context suggests a product, use it unless the conversation clearly contradicts it.
If no product specs match, return "None" and explain why the call doesn't clearly indicate a specific product.

Evaluate all parameters and return ONLY the JSON response.

Treat product accuracy as a strict QA test: include exact call-vs-spec comparison rows for any product or policy claim, and fail closed on unconfirmed facts.

Ensure param_comments contains exactly 10 entries in score order and each entry clearly explains the score using transcript evidence so it can be shown in hover tooltips."""

    try:
        response = await openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_msg}
            ],
            temperature=0.0,
            response_format={"type": "json_object"},
            timeout=120,
        )
        log_api_usage(OPENAI_MODEL, response.usage.prompt_tokens, response.usage.completion_tokens)
        raw = response.choices[0].message.content
        return json.loads(raw)
    except openai.AuthenticationError:
        return _mock_analysis(turns)
    except Exception as e:
        log.error(f"GPT-4o error: {e}")
        return _mock_analysis(turns)

def _mock_analysis(turns: List[Dict]) -> Dict:
    """Fallback mock analysis when API key not set"""
    import random
    random.seed(len(turns))
    
    scores = {
        "greeting_opening": random.randint(3, 5),
        "query_understanding": random.randint(2, 5),
        "response_accuracy": random.randint(2, 5),
        "communication_quality": random.randint(3, 5),
        "compliance": random.randint(3, 5),
        "personalisation": random.randint(2, 5),
        "empathy_soft_skills": random.randint(2, 5),
        "resolution": random.randint(2, 5),
        "system_behaviour": random.randint(3, 5),
        "closing_interaction": random.randint(2, 5)
    }
    
    weights = {"greeting_opening":5,"query_understanding":15,"response_accuracy":25,
               "communication_quality":10,"compliance":20,"personalisation":5,
               "empathy_soft_skills":5,"resolution":10,"system_behaviour":3,"closing_interaction":2}
    mins = {"greeting_opening":3,"query_understanding":3,"response_accuracy":4,
            "communication_quality":3,"compliance":4,"personalisation":3,
            "empathy_soft_skills":3,"resolution":3,"system_behaviour":3,"closing_interaction":3}
    
    weighted = sum(scores[k] * weights[k] / 5 for k in scores)
    failed = [k for k in scores if scores[k] < mins[k]]
    severity = "fatal" if scores["compliance"] < 3 else ("critical" if failed else ("watch" if weighted < 70 else "normal"))
    
    sentiments = ["positive","neutral","frustrated","angry","distressed"]
    
    return {
        "scores": scores,
        "param_comments": [
            "Polite opening that sets a professional tone." if scores["greeting_opening"] >= 4 else "Opening is abrupt or too generic.",
            "The reply matches the customer's question." if scores["query_understanding"] >= 4 else "The bot misses part of the request.",
            "No obvious factual error is visible in this mock flow." if scores["response_accuracy"] >= 4 else "The answer appears incomplete or unreliable.",
            "The explanation is easy to follow." if scores["communication_quality"] >= 4 else "The explanation is vague or poorly structured.",
            "Compliance appears acceptable in this mock flow." if scores["compliance"] >= 4 else "A likely compliance gap is present.",
            "The response uses available context naturally." if scores["personalisation"] >= 4 else "The reply is generic and not personalized.",
            "Warm and supportive tone throughout." if scores["empathy_soft_skills"] >= 4 else "The tone feels robotic or detached.",
            "The issue is resolved with a clear next step." if scores["resolution"] >= 4 else "The issue remains only partially resolved.",
            "No visible loop or system instability." if scores["system_behaviour"] >= 4 else "The flow shows repetition or instability.",
            "The call closes politely and cleanly." if scores["closing_interaction"] >= 4 else "The closing is short or missing.",
        ],
        "weighted_score": round(weighted, 2),
        "pass_fail": "FAIL" if failed else "PASS",
        "failed_parameters": failed,
        "category": random.choice(["Premium Receipt", "Premium Payment Assistance", "Policy Status Inquiry", "Policy Document Request", "Surrender Request", "Loan Against Policy", "Nominee Update", "Address or Contact Update", "Maturity Claim", "Death Claim", "Revival of Lapsed Policy", "Free Look Cancellation", "Rider Inquiry", "Fund Switch or Redirection", "Partial Withdrawal", "Benefit Illustration Request", "Complaint or Grievance", "Escalation Request", "General Inquiry"]),
        "severity": severity,
        "fatal_reason": "Compliance score below threshold" if severity == "fatal" else "",
        "flags": ["compliance_breach"] if severity == "fatal" else [],
        "sentiment": random.choice(sentiments),
        "product_mentioned": random.choice(["Bajaj Life Goal Suraksha","Bajaj Life Smart Wealth Goal","None"]),
        "product_confidence": round(random.uniform(0.48, 0.91), 2),
        "product_signals": ["premium", "policy term", "benefit"],
        "product_accuracy_score": random.randint(2, 5),
        "product_issues": "Demo mode – set OPENAI_API_KEY for real analysis",
        "what_should_have_been_said": "Demo mode – set OPENAI_API_KEY for real GPT-4o analysis",
        "strengths": "Demo mode analysis",
        "summary": "This is a demo analysis. Please set your OPENAI_API_KEY environment variable for real GPT-4o powered evaluations.",
        "product_checks": [],
        "turn_count": len(turns),
        "bot_turns": sum(1 for t in turns if t.get("speaker") in ["bot","agent","system"]),
        "customer_turns": sum(1 for t in turns if t.get("speaker") in ["customer","user"]),
        "estimated_duration_minutes": round(len(turns) * 0.5, 1)
    }
def _pick_best_product_name(
    explicitly_mentioned_set: set,
    gpt_product: str,
    rag_hit: Dict[str, Any],
) -> str:
    """Resolve the final product_mentioned value deterministically instead of
    using next(iter(set)) on an unordered set, which previously caused
    arbitrary/wrong product names to win when multiple candidates existed."""
    if not explicitly_mentioned_set:
        if gpt_product and gpt_product.lower() not in {"none", "unknown", ""}:
            return gpt_product
        rag_product = _safe_filename_label(str(rag_hit.get("product") or "None"))
        return rag_product if rag_product.lower() not in {"none", "unknown"} else "None"

    if gpt_product in explicitly_mentioned_set:
        return gpt_product

    if len(explicitly_mentioned_set) == 1:
        return next(iter(explicitly_mentioned_set))

    # Multiple distinct explicit candidates — rank by RAG score so the one
    # actually discussed (highest evidence) wins, not an arbitrary pick.
    all_scores = rag_hit.get("all_product_scores") or {}
    ranked = sorted(
        explicitly_mentioned_set,
        key=lambda name: all_scores.get(name, 0.0),
        reverse=True,
    )
    return ranked[0]
# ── Processing Pipeline ──────────────────────────────────────────────────────
async def process_realtime_call(job_id: str, item: dict):
    await asyncio.to_thread(update_job, job_id, status="processing")

    try:
        turns = item.get("turns") or parse_transcript_text(item["text"])
        meta  = item.get("meta", {})
        call_id = str(uuid.uuid4())

        # Store raw transcript immediately (Table 1) — off event loop
        await asyncio.to_thread(insert_raw_call, call_id, job_id, item)

        # SentenceTransformer inference — CPU-bound; run in thread
        rag_hit = await asyncio.to_thread(infer_product_context, item["text"])
        product_context = json.dumps(rag_hit, indent=2, ensure_ascii=False)
        detected_product = rag_hit.get("product", "None")
        if detected_product and detected_product != "None":
            product_context = f"Detected product guess: {detected_product}\n\n{product_context}"

        # GPT analysis — truly async
        analysis = await analyze_call_with_gpt4o(item["text"], turns, product_context)
        # Defensive normalization — 4o-mini is less consistent with types than 4o.
        if not isinstance(analysis.get("flags"), list):
            analysis["flags"] = [analysis["flags"]] if analysis.get("flags") else []
        if not isinstance(analysis.get("failed_parameters"), list):
            analysis["failed_parameters"] = []
        if not isinstance(analysis.get("scores"), dict):
            analysis["scores"] = {}
        if not isinstance(analysis.get("product_checks"), list):
            analysis["product_checks"] = []
 
        # explicitly_mentioned_set = set(
        #     _safe_filename_label(p)
        #     for p in (rag_hit.get("explicitly_mentioned") or [])
        # ) | set(
        #     _safe_filename_label(p)
        #     for p in _extract_product_mentions_from_text(item["text"] or "")
        # )
        # gpt_product = _safe_filename_label(str(analysis.get("product_mentioned") or "None"))
        # if gpt_product and gpt_product.lower() not in {"none", "unknown", ""}:
        #     analysis["product_mentioned"] = gpt_product
        # elif explicitly_mentioned_set:
        #     analysis["product_mentioned"] = next(iter(explicitly_mentioned_set))
        # else:
        #     rag_product = _safe_filename_label(str(rag_hit.get("product") or "None"))
        #     if rag_product and rag_product.lower() not in {"none", "unknown"} \
        #             and rag_product.lower() in (item["text"] or "").lower():
        #         analysis["product_mentioned"] = rag_product
        #     else:
        #         analysis["product_mentioned"] = "None"
        explicitly_mentioned_set = set(
            _safe_filename_label(p) for p in (rag_hit.get("explicitly_mentioned") or [])
        ) | set(
            _safe_filename_label(p) for p in _extract_product_mentions_from_text(item["text"] or "")
        )
        gpt_product = _safe_filename_label(str(analysis.get("product_mentioned") or "None"))

        # if explicitly_mentioned_set:
        #     # Transcript explicitly named a product — trust this over the model's guess,
        #     # but prefer the model's pick IF it's actually inside the explicit set.
        #     if gpt_product in explicitly_mentioned_set:
        #         analysis["product_mentioned"] = gpt_product
        #     else:
        #         analysis["product_mentioned"] = next(iter(explicitly_mentioned_set))
        # elif gpt_product and gpt_product.lower() not in {"none", "unknown", ""}:
        #     analysis["product_mentioned"] = gpt_product
        # else:
        #     rag_product = _safe_filename_label(str(rag_hit.get("product") or "None"))
        #     if rag_product and rag_product.lower() not in {"none", "unknown"} \
        #             and rag_product.lower() in (item["text"] or "").lower():
        #         analysis["product_mentioned"] = rag_product
        #     else:
        #         analysis["product_mentioned"] = "None"
        analysis["product_mentioned"] = _pick_best_product_name(
            explicitly_mentioned_set, gpt_product, rag_hit
        )
        
        analysis["products_mentioned"] = _collect_products_mentioned(analysis, rag_hit, item["text"])
        analysis.setdefault("product_confidence", rag_hit.get("confidence", 0.0))
        analysis.setdefault("product_signals", rag_hit.get("matched_terms", []))
        analysis.setdefault("product_checks", [])
        analysis.setdefault("product_profile", rag_hit.get("profile", {}))
        analysis.setdefault("product_evidence", rag_hit.get("chunks", []))
        analysis.setdefault("secondary_products", rag_hit.get("secondary_products", []))
        analysis.setdefault("all_product_scores", rag_hit.get("all_product_scores", {}))
        analysis["sentiment"] = _refine_sentiment(turns, analysis.get("sentiment", "neutral"))
        if not analysis.get("param_comments") or len(analysis.get("param_comments", [])) < len(PARAM_ORDER):
            analysis["param_comments"] = _fallback_param_comments(analysis.get("scores", {}))
        analysis = _apply_qa_policy_rules(analysis, turns)
        analysis["annotated_transcript"] = _annotate_transcript(turns, analysis)
        # Always recompute score server-side
        analysis["weighted_score"] = _compute_weighted_score(analysis.get("scores", {}))
 
        # Overlay voicebot metadata
        if meta.get("intent"):         analysis["voicebot_intent"] = meta["intent"]
        if meta.get("summary","").strip(): analysis["summary"] = meta["summary"].strip()
        if meta.get("no_of_queries") is not None:          analysis["no_of_queries"] = meta["no_of_queries"]
        if meta.get("no_of_queries_resolved") is not None: analysis["no_of_queries_resolved"] = meta["no_of_queries_resolved"]
        if meta.get("customer_journey"): analysis["customer_journey"] = meta["customer_journey"]
        if meta.get("standalone_type"):
            analysis["category"] = meta["standalone_type"][0] if meta["standalone_type"] else analysis.get("category","General Query")
            analysis["standalone_type"] = meta["standalone_type"]
        if meta.get("sub_query_type"): analysis["sub_query_type"] = meta["sub_query_type"]
        if meta.get("call_comp_flag"): analysis["call_comp_flag"] = meta["call_comp_flag"]
        if meta.get("audio_link"):     analysis["audio_link"] = meta["audio_link"]
 
        # ── Derive fatal/flagged deterministically ────────────────────────────
        is_fatal   = analysis.get("severity") in {"fatal", "critical"}
        is_flagged = len(analysis.get("flags", [])) > 0
 
        call_record = {
            "id": call_id,
            "job_id": job_id,
            "name": item["name"],
            "sl": item.get("sl", ""),
            "source_file": "realtime",
            "transcript": turns,
            "raw_text": item["text"][:5000],
            "analysis": analysis,
            "processed_at": datetime.now().isoformat(),
            "flagged": is_flagged,
            "fatal":   is_fatal,
        }
 
        # Store analysis result (Table 2) — off event loop
        await asyncio.to_thread(upsert_analyzed_call, call_record)

        await asyncio.to_thread(update_job, job_id,
                   status="completed",
                   processed=1,
                   fatal_count=1 if is_fatal else 0,
                   flag_count=1 if is_flagged else 0,
                   completed_at=datetime.now().isoformat())

    except Exception as e:
        log.error(f"Realtime call processing error: {e}", exc_info=True)
        await asyncio.to_thread(update_job, job_id, status="completed", completed_at=datetime.now().isoformat())
 

async def process_job(job_id: str, file_paths: List[Path]):
    job = await asyncio.to_thread(get_job, job_id)
    if not job:
        return

    all_transcripts = []
    for fp in file_paths:
        try:
            items = await asyncio.to_thread(extract_transcripts_from_file, fp)
            all_transcripts.extend(items)
        except Exception as e:
            log.error(f"File parse error {fp}: {e}")

    await asyncio.to_thread(update_job, job_id, total=len(all_transcripts), status="processing", processed=0)

    for item in all_transcripts:
        signal = check_job_signal(job_id)

        if signal == "cancel":
            clear_job_signal(job_id)
            await asyncio.to_thread(update_job, job_id, status="cancelled", completed_at=datetime.now().isoformat())
            log.info(f"Job {job_id} cancelled")
            return

        if signal == "pause":
            clear_job_signal(job_id)
            await asyncio.to_thread(update_job, job_id, status="paused")
            log.info(f"Job {job_id} paused")
            while True:
                await asyncio.sleep(2)
                wake = check_job_signal(job_id)
                if wake == "cancel":
                    clear_job_signal(job_id)
                    await asyncio.to_thread(update_job, job_id, status="cancelled", completed_at=datetime.now().isoformat())
                    return
                if wake == "resume":
                    clear_job_signal(job_id)
                    await asyncio.to_thread(update_job, job_id, status="processing")
                    log.info(f"Job {job_id} resumed")
                    break

        try:
            call_id = str(uuid.uuid4())
            turns   = parse_transcript_text(item["text"])

            # Store raw transcript (Table 1) — off event loop
            await asyncio.to_thread(insert_raw_call, call_id, job_id, item)

            # Product context — SentenceTransformer inference is CPU-bound; run in thread
            rag_hit = await asyncio.to_thread(infer_product_context, item["text"])
            product_context = json.dumps(rag_hit, indent=2, ensure_ascii=False)
            detected_product = rag_hit.get("product", "None")
            if detected_product and detected_product != "None":
                product_context = f"Detected product guess: {detected_product}\n\n{product_context}"

            # GPT analysis — truly async now (AsyncAzureOpenAI / AsyncOpenAI)
            analysis = await analyze_call_with_gpt4o(item["text"], turns, product_context)
            # Defensive normalization
            if not isinstance(analysis.get("flags"), list):
                analysis["flags"] = [analysis["flags"]] if analysis.get("flags") else []
            if not isinstance(analysis.get("failed_parameters"), list):
                analysis["failed_parameters"] = []
            if not isinstance(analysis.get("scores"), dict):
                analysis["scores"] = {}
            if not isinstance(analysis.get("product_checks"), list):
                analysis["product_checks"] = []

            explicitly_mentioned_set = set(
                _safe_filename_label(p) for p in (rag_hit.get("explicitly_mentioned") or [])
            ) | set(
                _safe_filename_label(p) for p in _extract_product_mentions_from_text(item["text"] or "")
            )
            gpt_product = _safe_filename_label(str(analysis.get("product_mentioned") or "None"))

            analysis["product_mentioned"] = _pick_best_product_name(
                explicitly_mentioned_set, gpt_product, rag_hit
            )

            analysis["products_mentioned"] = _collect_products_mentioned(analysis, rag_hit, item["text"])
            analysis.setdefault("product_confidence", rag_hit.get("confidence", 0.0))
            analysis.setdefault("product_signals", rag_hit.get("matched_terms", []))
            analysis.setdefault("product_checks", [])
            analysis.setdefault("product_profile", rag_hit.get("profile", {}))
            analysis.setdefault("product_evidence", rag_hit.get("chunks", []))
            analysis.setdefault("secondary_products", rag_hit.get("secondary_products", []))
            analysis.setdefault("all_product_scores", rag_hit.get("all_product_scores", {}))
            analysis["sentiment"] = _refine_sentiment(turns, analysis.get("sentiment", "neutral"))
            if not analysis.get("param_comments") or len(analysis.get("param_comments", [])) < len(PARAM_ORDER):
                analysis["param_comments"] = _fallback_param_comments(analysis.get("scores", {}))
            analysis = _apply_qa_policy_rules(analysis, turns)
            analysis["annotated_transcript"] = _annotate_transcript(turns, analysis)
            analysis["weighted_score"] = _compute_weighted_score(analysis.get("scores", {}))

            is_fatal   = analysis.get("severity") in {"fatal", "critical"}
            is_flagged = len(analysis.get("flags", [])) > 0

            call_record = {
                "id": call_id,
                "job_id": job_id,
                "name": item["name"],
                "sl": item.get("sl", ""),
                "source_file": item.get("source_file", ""),
                "transcript": turns if turns else [{"sl":1,"speaker":"unknown","text":item["text"][:2000]}],
                "raw_text": item["text"][:5000],
                "analysis": analysis,
                "processed_at": datetime.now().isoformat(),
                "flagged": is_flagged,
                "fatal":   is_fatal,
            }

            # Store analysis result — off event loop
            await asyncio.to_thread(upsert_analyzed_call, call_record)
            await asyncio.to_thread(increment_job_counters, job_id,
                                    1, 1 if is_fatal else 0, 1 if is_flagged else 0)

        except Exception as e:
            log.error(f"Call processing error [{item.get('name','?')}]: {e}", exc_info=True)
            # still count as processed so progress bar advances
            await asyncio.to_thread(increment_job_counters, job_id, 1, 0, 0)

    job_now = await asyncio.to_thread(get_job, job_id)
    if job_now and job_now.get("status") not in {"cancelled", "paused"}:
        await asyncio.to_thread(update_job, job_id, status="completed", completed_at=datetime.now().isoformat())
 
# ── API Endpoints ────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    return FileResponse(str(FRONTEND_DIR / "templates" / "index.html"))

@app.head("/")
async def head_root():
    """Render health-check uses HEAD / — must return 200 not 405."""
    return JSONResponse(content={}, status_code=200)

@app.post("/api/upload")
async def upload_files(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "No files provided")

    ALLOWED = {".xlsx", ".xls", ".pdf", ".docx", ".doc", ".txt", ".json"}
    for f in files:
        ext = Path(f.filename).suffix.lower()
        if ext not in ALLOWED:
            raise HTTPException(400, f"Unsupported file type: {f.filename}. Allowed: {', '.join(sorted(ALLOWED))}")

    job_id = str(uuid.uuid4())
    job_dir = UPLOAD_DIR / job_id
    job_dir.mkdir(parents=True)

    saved_paths = []
    for f in files:
        dest = job_dir / f.filename
        content = await f.read()
        dest.write_bytes(content)
        saved_paths.append(dest)
    
    # db = load_db()
    # db["jobs"].append({
    #     "id": job_id,
    #     "status": "queued",
    #     "files": [f.filename for f in files],
    #     "total": 0,
    #     "processed": 0,
    #     "fatal_count": 0,
    #     "flag_count": 0,
    #     "created_at": datetime.now().isoformat(),
    #     "completed_at": None
    # })
    # save_db(db)
    create_job({
        "id": job_id,
        "status": "queued",
        "files": [f.filename for f in files],
        "total": 0,
        "processed": 0,
        "fatal_count": 0,
        "flag_count": 0,
        "created_at": datetime.now().isoformat(),
        "completed_at": None,
    })
    
    background_tasks.add_task(process_job, job_id, saved_paths)
    return {"job_id": job_id, "files_received": len(files), "status": "queued"}
from fastapi import Request

def _build_ingest_item(payload: dict) -> dict:
    """Build a normalised item dict from a single call payload (vendor or extracted_logs format)."""
    role_map = {"assistant": "bot", "user": "customer"}
    # Support both unique_call_id (vendor) and sender_id (extracted_logs)
    call_name = (
        payload.get("unique_call_id")
        or payload.get("sender_id")
        or str(uuid.uuid4())
    )
    turns = []
    for i, msg in enumerate(payload.get("conversation_log", []), start=1):
        if not isinstance(msg, dict):
            continue
        content = (msg.get("content") or "").strip()
        if not content:
            continue
        turns.append({
            "sl": i,
            "speaker": role_map.get(msg.get("role", "user"), "unknown"),
            "text": content,
        })
    raw_text = "\n".join(f"{t['speaker']}: {t['text']}" for t in turns)
    return {
        "name": str(call_name),
        "text": raw_text,
        "sl": str(call_name),
        "source_file": "realtime",
        "turns": turns,
        "meta": {
            "intent": payload.get("INTENT"),
            "stage_code": payload.get("STAGE_CODE"),
            "sentiment": payload.get("sentiment"),
            "summary": payload.get("conversation_summary"),
            "no_of_queries": payload.get("no_of_queries"),
            "no_of_queries_resolved": payload.get("no_of_queries_resolved"),
            "call_comp_flag": payload.get("call_comp_flag"),
            "customer_journey": payload.get("customer_journey", []),
            "standalone_type": payload.get("standalone_type", []),
            "sub_query_type": payload.get("sub_query_type", []),
            "root_dscn": payload.get("root_dscn"),
            "mid_dscn": payload.get("mid_dscn"),
            "root_tta": payload.get("root_tta"),
            "mid_tta": payload.get("mid_tta"),
            "audio_link": payload.get("audio_link") or "",
        },
    }

@app.post("/api/ingest-call")
async def ingest_realtime_call(request: Request, background_tasks: BackgroundTasks):
    """Receive one or many call payloads in the voicebot API format.

    Accepts:
      - Single call object  : { "unique_call_id": "...", "conversation_log": [...], ... }
      - Array of call objects: [ { "sender_id": "...", "conversation_log": [...] }, ... ]

    Both vendor single-call pushes and bulk extracted_logs arrays are supported.
    """
    # Auth check (Bearer token or X-API-Key)
    expected_key = _env("INGEST_API_KEY", "")
    if expected_key:
        auth = request.headers.get("authorization", "") or request.headers.get("Authorization", "")
        api_key_header = request.headers.get("x-api-key", "") or request.headers.get("X-API-Key", "")
        provided = None
        if auth and auth.lower().startswith("bearer "):
            provided = auth.split(" ", 1)[1].strip()
        elif api_key_header:
            provided = api_key_header.strip()
        if not provided or provided != expected_key:
            raise HTTPException(status_code=401, detail="Unauthorized")

    body = await request.json()

    # ── Single call (vendor default) ────────────────────────────────────────
    if isinstance(body, dict):
        item = _build_ingest_item(body)
        if not item["turns"]:
            raise HTTPException(status_code=422, detail="conversation_log has no valid messages")
        job_id = str(uuid.uuid4())
        # db = load_db()
        # db["jobs"].append({
        #     "id": job_id, "status": "queued",
        #     "files": [f"realtime:{item['name']}"],
        #     "total": 1, "processed": 0,
        #     "fatal_count": 0, "flag_count": 0,
        #     "created_at": datetime.now().isoformat(), "completed_at": None,
        # })
        # save_db(db)
        create_job({
            "id": job_id, "status": "queued",
            "files": [f"realtime:{item['name']}"],
            "total": 1, "processed": 0,
            "fatal_count": 0, "flag_count": 0,
            "created_at": datetime.now().isoformat(), "completed_at": None,
        })
        background_tasks.add_task(process_realtime_call, job_id, item)
        return {"job_id": job_id, "call_id": item["name"], "status": "queued"}

    # ── Array of calls (extracted_logs / bulk) ───────────────────────────────
    if isinstance(body, list):
        if not body:
            raise HTTPException(status_code=422, detail="Empty array")
        results, skipped = [], []
        for record in body:
            if not isinstance(record, dict):
                skipped.append({"reason": "not an object"})
                continue
            item = _build_ingest_item(record)
            if not item["turns"]:
                skipped.append({"call_id": item["name"], "reason": "no valid messages"})
                continue
            job_id = str(uuid.uuid4())
            # db = load_db()
            # db["jobs"].append({
            #     "id": job_id, "status": "queued",
            #     "files": [f"realtime:{item['name']}"],
            #     "total": 1, "processed": 0,
            #     "fatal_count": 0, "flag_count": 0,
            #     "created_at": datetime.now().isoformat(), "completed_at": None,
            # })
            # save_db(db)
            create_job({
                "id": job_id, "status": "queued",
                "files": [f"realtime:{item['name']}"],
                "total": 1, "processed": 0,
                "fatal_count": 0, "flag_count": 0,
                "created_at": datetime.now().isoformat(), "completed_at": None,
            })
            background_tasks.add_task(process_realtime_call, job_id, item)
            results.append({"job_id": job_id, "call_id": item["name"], "status": "queued"})
        return {
            "queued": len(results),
            "skipped": len(skipped),
            "results": results,
            "skipped_details": skipped,
        }

    raise HTTPException(status_code=422, detail="Payload must be a call object or array of call objects")

@app.post("/api/upload-products")
async def upload_product_specs(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    indexed = []
    for f in files:
        if f.filename.endswith(".pdf"):
            dest = PRODUCT_DIR / f.filename
            content = await f.read()
            dest.write_bytes(content)
            indexed.append(f.filename)
    background_tasks.add_task(rebuild_product_rag_index)
    mode = json.loads(RAG_BACKEND_FILE.read_text()).get("mode", "faiss") if RAG_BACKEND_FILE.exists() else "faiss"
    message = f"Rebuilt product knowledge base after indexing {len(indexed)} PDF(s)"
    return {"indexed": indexed, "mode": mode, "message": message}

@app.get("/api/product-specs")
async def list_product_specs():
    meta_rows = _load_rag_meta()
    chunk_map = Counter(row.get("source", "") for row in meta_rows)

    specs = []
    for p in sorted(PRODUCT_DIR.glob("*.pdf")):
        stat = p.stat()
        specs.append({
            "name": p.name,
            "display_name": _safe_filename_label(p.name),
            "size": stat.st_size,
            "updated_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "chunks": chunk_map.get(p.name, 0)
        })

    return {
        "mode": json.loads(RAG_BACKEND_FILE.read_text()).get("mode", "faiss") if RAG_BACKEND_FILE.exists() else "faiss",
        "count": len(specs),
        "specs": specs
    }

@app.delete("/api/product-specs/{filename}")
async def delete_product_spec(filename: str):
    safe_name = Path(filename).name
    target = PRODUCT_DIR / safe_name

    if not target.exists():
        raise HTTPException(404, "Product spec not found")

    try:
        target.unlink()
    except Exception as e:
        raise HTTPException(500, f"Failed to delete file: {e}")

    rebuild_product_rag_index()

    return {
        "deleted": safe_name,
        "removed_from": ["file", "rag"],
        "message": f"Removed {safe_name}"
    }

@app.delete("/api/product-specs")
async def delete_all_product_specs():
    remove_all_product_specs()
    return {"message": "All product specs removed"}

@app.get("/api/jobs")
async def get_jobs():
    return list_jobs()
 
@app.get("/api/jobs/{job_id}")
async def get_job_endpoint(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    return job


@app.get("/api/calls")
async def get_calls(
    page: int = Query(1, ge=1),
    page_size: int = Query(50, ge=1, le=500),
    severity: Optional[str] = None,
    category: Optional[str] = None,
    sentiment: Optional[str] = None,
    pass_fail: Optional[str] = None,
    flagged: Optional[bool] = None,
    search: Optional[str] = None,
    job_id: Optional[str] = None,
    sort_by: str = "processed_at",
    sort_dir: str = "desc",
):
    return list_analyzed_calls(
        page=page, page_size=page_size,
        severity=severity, category=category, sentiment=sentiment,
        pass_fail=pass_fail, flagged=flagged, job_id=job_id,
        search=search, sort_by=sort_by, sort_dir=sort_dir,
    )


# ── /api/calls/{call_id} ─────────────────────────────────────────────────────
@app.get("/api/calls/{call_id}")
async def get_call_detail(call_id: str):
    call = get_analyzed_call(call_id)
    if not call:
        raise HTTPException(404, "Call not found")
    a = call.get("analysis", {})
    if not a.get("param_comments") or len(a.get("param_comments", [])) < len(PARAM_ORDER):
        a["param_comments"] = _fallback_param_comments(a.get("scores", {}))
    if not a.get("annotated_transcript"):
        a["annotated_transcript"] = _annotate_transcript(
            call.get("transcript", []), a
        )
    call["analysis"] = a
    return call


@app.get("/api/export/calls.xlsx")
async def export_calls_excel():
    from db import db_cursor  # already imported but explicit here for clarity
    with db_cursor(commit=False) as cur:
        cur.execute("SELECT * FROM analyzed_calls ORDER BY processed_at DESC")
        rows = cur.fetchall()
 
    export_rows = []
    for r in rows:
        a = r.get("analysis") or {}
        scores = a.get("scores") or {}
        export_rows.append({
            "id":                    r.get("id",""),
            "job_id":                r.get("job_id",""),
            "name":                  r.get("name",""),
            "sl":                    r.get("sl",""),
            "processed_at":          str(r.get("processed_at","")),
            "pass_fail":             r.get("pass_fail",""),
            "severity":              r.get("severity",""),
            "category":              r.get("category",""),
            "sentiment":             r.get("sentiment",""),
            "weighted_score":        float(r.get("weighted_score") or 0),
            "failed_parameters":     ", ".join(r.get("failed_parameters") or []),
            "flags":                 ", ".join(r.get("flags") or []),
            "fatal":                 r.get("fatal", False),
            "flagged":               r.get("flagged", False),
            "product_mentioned":     r.get("product_mentioned",""),
            "product_confidence":    float(r.get("product_confidence") or 0),
            "product_accuracy_score":r.get("product_accuracy_score",""),
            "product_issues":        r.get("product_issues",""),
            "strengths":             a.get("strengths",""),
            "summary":               a.get("summary",""),
            "greeting_opening":      scores.get("greeting_opening",""),
            "query_understanding":   scores.get("query_understanding",""),
            "response_accuracy":     scores.get("response_accuracy",""),
            "communication_quality": scores.get("communication_quality",""),
            "compliance":            scores.get("compliance",""),
            "personalisation":       scores.get("personalisation",""),
            "empathy_soft_skills":   scores.get("empathy_soft_skills",""),
            "resolution":            scores.get("resolution",""),
            "system_behaviour":      scores.get("system_behaviour",""),
            "closing_interaction":   scores.get("closing_interaction",""),
        })
 
    df = pd.DataFrame(export_rows)
    stream = io.BytesIO()
    with pd.ExcelWriter(stream, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="calls")
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=calls_export.xlsx"},
    )

@app.get("/api/calls/{call_id}/report.pdf")
async def export_call_report_pdf(call_id: str):
    db = load_db()
    call = next((c for c in db.get("calls", []) if c.get("id") == call_id), None)
    if not call:
        raise HTTPException(404, "Call not found")

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except Exception:
        raise HTTPException(500, "PDF export requires reportlab. Install with pip install reportlab")

    a = call.get("analysis", {})
    scores = a.get("scores", {})
    stream = io.BytesIO()
    doc = SimpleDocTemplate(stream, pagesize=A4, title=f"Call Report {call.get('name', '')}")
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Bajaj Life Insurance - Call QA Report", styles["Title"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"Call: {call.get('name', '')}", styles["Heading3"]))
    story.append(Paragraph(f"Processed At: {call.get('processed_at', '')}", styles["Normal"]))
    story.append(Paragraph(f"Score: {a.get('weighted_score', 0)} | Status: {a.get('pass_fail', '')} | Severity: {a.get('severity', '')}", styles["Normal"]))
    story.append(Paragraph(f"Product: {a.get('product_mentioned', 'None')} (confidence: {a.get('product_confidence', 0)})", styles["Normal"]))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Summary", styles["Heading3"]))
    story.append(Paragraph(a.get("summary", ""), styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Parameter Scores", styles["Heading3"]))
    for key in PARAM_ORDER:
        story.append(Paragraph(f"{key.replace('_', ' ').title()}: {scores.get(key, '-')}/5", styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("QA Findings", styles["Heading3"]))
    for finding in a.get("qa_findings", []):
        story.append(Paragraph(f"- [{finding.get('type', 'info')}] {finding.get('text', '')}", styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Product Checks", styles["Heading3"]))
    for pc in a.get("product_checks", [])[:6]:
        story.append(Paragraph(f"- {pc.get('vtext', pc.get('verdict', 'check'))}: {pc.get('stmt', '')}", styles["Normal"]))
        story.append(Paragraph(f"  Fact: {pc.get('fact', '')}", styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Annotated Transcript", styles["Heading3"]))
    turns = a.get("annotated_transcript") or call.get("transcript", [])
    for turn in turns[:120]:
        speaker = (turn.get("speaker") or "unknown").upper()
        text = turn.get("text", "")
        tags = turn.get("tags", [])
        tag_txt = f" [{' | '.join(tags)}]" if tags else ""
        story.append(Paragraph(f"{turn.get('sl', '')}. {speaker}: {text}{tag_txt}", styles["Normal"]))

    doc.build(story)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=call_report_{call_id}.pdf"},
    )


@app.delete("/api/calls/{call_id}")
async def delete_call(call_id: str):
    deleted = delete_analyzed_call(call_id)
    if not deleted:
        raise HTTPException(404, "Call not found")
    return {"deleted": call_id, "message": "Call removed"}


@app.get("/api/dashboard")
async def get_dashboard():
    return get_dashboard_stats()

@app.post("/api/jobs/{job_id}/cancel")
async def cancel_job(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    if job.get("status") in {"completed", "cancelled", "failed"}:
        raise HTTPException(400, f"Job is already {job['status']}")
    signal_job(job_id, "cancel")
    return {"job_id": job_id, "signal": "cancel",
            "message": "Cancellation signal sent — will stop after current transcript"}
@app.post("/api/jobs/{job_id}/pause")
async def pause_job(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    if job.get("status") != "processing":
        raise HTTPException(400, f"Can only pause a processing job (current: {job.get('status')})")
    signal_job(job_id, "pause")
    return {"job_id": job_id, "signal": "pause",
            "message": "Pause signal sent — will pause after current transcript"}
 
@app.post("/api/jobs/{job_id}/resume")
async def resume_job(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    if job.get("status") != "paused":
        raise HTTPException(400, f"Can only resume a paused job (current: {job.get('status')})")
    signal_job(job_id, "resume")
    return {"job_id": job_id, "signal": "resume", "message": "Resume signal sent"}

@app.get("/api/fatal-calls")
async def get_fatal_calls_endpoint():
    return db_get_fatal_calls()

@app.delete("/api/calls")
async def clear_all_calls():
    clear_all_data()
    for job_dir in UPLOAD_DIR.glob("*"):
        if job_dir.is_dir():
            shutil.rmtree(job_dir, ignore_errors=True)
    return {"message": "All data cleared"}
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
